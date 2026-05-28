"""
Outgoing registry API router.
"""
import base64
import html
import json
import re
import subprocess
import tempfile
import unicodedata
import uuid
from datetime import date, datetime, timedelta
from functools import lru_cache
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import List, Optional
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form, Request, Body, Response
from sqlalchemy import select, func, or_, and_, delete
from sqlalchemy.exc import IntegrityError, OperationalError
from sqlalchemy.ext.asyncio import AsyncSession
from starlette.concurrency import run_in_threadpool

from app.core.config import settings
from app.database.session import get_db
from app.models import (
    Company,
    Contract,
    ContractDocument,
    Deal,
    IncomeExpenseEntry,
    OutgoingDocument,
    OutgoingDocumentVersion,
    OutgoingDocumentFile,
    OutgoingNumberSequence,
    OutgoingDailyNumberSequence,
    Document,
    DocumentTemplate,
    DocumentTemplateVersion,
    Stage,
)
from app.schemas.outgoing_document import (
    OutgoingDocumentResponse,
    OutgoingDocumentDetailResponse,
    OutgoingDocumentResolveRequest,
    OutgoingDocumentUpdate,
    OutgoingDocumentVersionResponse,
    OutgoingDocumentFileResponse,
)
from app.services.storage import (
    clean_name,
    ensure_path,
    publish,
    upload_bytes_with_safe_extension,
    delete_path,
    read_file_bytes,
    storage_available,
)
from app.core.auth_middleware import CurrentUser
from app.services.data_health import safe_refresh_deal_health_issues, safe_refresh_orphan_health_issues
from app.services.document_template_fields import get_template_fields
from app.services.sequence_lock import sequence_lock
from app.services.outgoing_document_editor import (
    EDITOR_SCHEMA_VERSION,
    default_editor_draft,
    get_editor_blocks,
    get_first_editor_block,
    get_editor_block_catalog,
    normalize_editor_draft,
    normalize_editor_mode,
    validate_editor_draft,
)
from app.services.permissions import allowed_deal_ids, get_section_permissions, ensure_can_edit_record
from app.services.event_log import log_event
from app.services.approval_runtime import ensure_entity_action_allowed


router = APIRouter()

COMPANY_OPTIONS = {
    "normbud": {"label": "\u041d\u041e\u0420\u041c\u0411\u0423\u0414"},
    "bayer": {"label": "\u0411\u0410\u0419\u0415\u0420"},
    "morozov": {"label": "\u0418\u041f \u041c\u043e\u0440\u043e\u0437\u043e\u0432"},
}
FORMAL_COMPANY_NAMES = {
    "normbud": "ООО «НОРМБУД»",
    "bayer": "ООО «БАЙЕР»",
    "morozov": "ИП Морозов О.А.",
}
FORMAL_COMPANY_ALIASES = {
    "normbud": ["нормбуд", "ооо нормбуд"],
    "bayer": ["байер", "ооо байер"],
    "morozov": ["ип морозов", "морозов о а", "морозов"],
}
OUR_COMPANY_REQUISITES = {
    "normbud": {
        "name": "ООО «НОРМБУД»",
        "short_name": "ООО «НОРМБУД»",
        "full_name": "Общество с ограниченной ответственностью «НОРМБУД»",
        "inn": "7733316255",
        "kpp": "771701001",
        "address": "129226, г. Москва, вн. тер. г. муниципальный округ Ростокино, ул. Сельскохозяйственная, д. 4 стр. 16",
    },
    "bayer": {
        "name": "ООО «БАЙЕР»",
        "short_name": "ООО «БАЙЕР»",
        "full_name": "Общество с ограниченной ответственностью «БАЙЕР»",
        "inn": "7722685412",
        "kpp": "771301001",
        "address": "127434, г. Москва, вн. тер. г. муниципальный округ Тимирязевский, ш. Дмитровское, д. 7 к. 2, помещ. 5А/1",
    },
    "morozov": {
        "name": "ИП Морозов О.А.",
        "short_name": "ИП Морозов О.А.",
        "full_name": "Индивидуальный предприниматель Морозов Олег Артурович",
        "inn": "",
        "kpp": "",
        "address": "129226, г. Москва, пр-т Мира, д. 179а, кв. 27",
    },
}
DOCUMENT_KIND_LABELS = {
    "letter": "Письмо",
    "invoice": "Счет",
    "upd": "УПД",
    "act": "Акт",
    "vat_invoice": "Счет-фактура",
}
DOCUMENT_REGISTRY_TYPES = {
    "letter": "outgoing_letter",
    "invoice": "outgoing_invoice",
    "upd": "outgoing_upd",
    "act": "outgoing_act",
    "vat_invoice": "outgoing_vat_invoice",
}
DOCUMENT_KIND_DEFAULT_SUBJECTS = {
    "letter": "Исходящее письмо",
    "invoice": "Счет",
    "upd": "УПД",
    "act": "Акт",
    "vat_invoice": "Счет-фактура",
}
FINANCIAL_DAILY_NUMBER_KINDS = {"invoice", "upd", "vat_invoice"}

PROJECT_ROOT = Path(__file__).resolve().parents[3]
FRONTEND_DIR = PROJECT_ROOT / "frontend"
RENDER_SCRIPT_PATH = FRONTEND_DIR / "scripts" / "render_outgoing_docx.mjs"
PREVIEW_ASSETS_DIR = PROJECT_ROOT / "frontend" / "public" / "templates" / "_extracted_preview"
COMPANY_RENDER_PROFILES = {
    "bayer": {
        "logo": "outgoing_bayer_image2.png",
        "signature": "outgoing_bayer_image1.png",
        "info_lines": [
            "Общество с ограниченной ответственностью «БАЙЕР»",
            "(ООО «БАЙЕР»)",
            "Эл. почта: info@byer.ru",
            "Тел. +7 (495) 128-11-77",
        ],
        "footer_lines": [
            "127434, г. Москва, вн. тер. г. муниципальный округ Тимирязевский, ш. Дмитровское, д. 7 к. 2, помещ. 5А/1",
            "ИНН 7722685412, КПП 771301001, ОГРН 1097746256211",
        ],
        "signer_title": "Генеральный директор",
        "signer_name": "О.А. Морозов",
    },
    "morozov": {
        "logo": "",
        "signature": "outgoing_morozov_image1.png",
        "info_lines": [
            "Индивидуальный предприниматель Морозов Олег Артурович",
            "ОГРНИП 318774600197203",
            "129226, г. Москва, пр-т Мира, д. 179а, кв. 27",
            "Эл. почта: morozov@proriski.ru",
            "Тел. +7 (903) 160-18-01",
        ],
        "footer_lines": [],
        "signer_title": "С уважением, Руководитель",
        "signer_name": "О.А. Морозов",
    },
    "normbud": {
        "logo": "outgoing_normbud_image2.png",
        "signature": "outgoing_normbud_image1.png",
        "info_lines": [
            "Общество с ограниченной ответственностью «НОРМБУД»",
            "(ООО «НОРМБУД»)",
            "129226, г. Москва, вн. тер. г. муниципальный округ Ростокино, ул. Сельскохозяйственная, д. 4 стр. 16",
            "ИНН 7733316255, КПП 771701001, ОГРН 1177746139636, e-mail: info@normbud.ru",
        ],
        "footer_lines": [],
        "signer_title": "Генеральный директор",
        "signer_name": "С.В. Воронин",
    },
}

if settings.APP_VARIANT == "test_portal":
    COMPANY_OPTIONS = {
        "normbud": {"label": "Nexus Beta"},
        "bayer": {"label": "Nexus Alpha"},
        "morozov": {"label": "Nexus Solo"},
    }
    COMPANY_RENDER_PROFILES = {
        "bayer": {
            "logo": "outgoing_bayer_image2.png",
            "signature": "outgoing_bayer_image1.png",
            "info_lines": [
                "Общество с ограниченной ответственностью «НЕКСУС АЛЬФА»",
                "(ООО «НЕКСУС АЛЬФА»)",
                "Эл. почта: alpha@nexus.test",
                "Тел. +7 (495) 000-10-10",
            ],
            "footer_lines": [
                "г. Москва, Тестовый проезд, д. 10",
                "ИНН 7700100010, КПП 770001001, ОГРН 1267700001010",
            ],
            "signer_title": "Генеральный директор",
            "signer_name": "Тест Тестов",
        },
        "morozov": {
            "logo": "",
            "signature": "outgoing_morozov_image1.png",
            "info_lines": [
                "Индивидуальный предприниматель Нексус Соло",
                "ОГРНИП 326770000303030",
                "г. Москва, Тестовый проезд, д. 30",
                "Эл. почта: solo@nexus.test",
                "Тел. +7 (495) 000-30-30",
            ],
            "footer_lines": [],
            "signer_title": "Руководитель",
            "signer_name": "Н. Соло",
        },
        "normbud": {
            "logo": "outgoing_normbud_image2.png",
            "signature": "outgoing_normbud_image1.png",
            "info_lines": [
                "Общество с ограниченной ответственностью «НЕКСУС БЕТА»",
                "(ООО «НЕКСУС БЕТА»)",
                "г. Москва, Тестовый проезд, д. 20",
                "ИНН 7700200020, КПП 770002001, ОГРН 1267700002020",
                "e-mail: beta@nexus.test",
            ],
            "footer_lines": [],
            "signer_title": "Генеральный директор",
            "signer_name": "Б. Тестова",
        },
    }
    OUR_COMPANY_REQUISITES = {
        "bayer": {
            "name": "Nexus Alpha",
            "short_name": "Nexus Alpha",
            "full_name": "ООО «Nexus Alpha»",
            "inn": "7700100010",
            "kpp": "770001001",
            "address": "г. Москва, Тестовый проезд, д. 10",
        },
        "morozov": {
            "name": "Nexus Solo",
            "short_name": "Nexus Solo",
            "full_name": "ИП Nexus Solo",
            "inn": "7700300030",
            "kpp": "",
            "address": "г. Москва, Тестовый проезд, д. 30",
        },
        "normbud": {
            "name": "Nexus Beta",
            "short_name": "Nexus Beta",
            "full_name": "ООО «Nexus Beta»",
            "inn": "7700200020",
            "kpp": "770002001",
            "address": "г. Москва, Тестовый проезд, д. 20",
        },
    }


def _company_label(company_key: Optional[str]) -> str:
    key = _normalize_company_key(company_key)
    return COMPANY_OPTIONS.get(key, COMPANY_OPTIONS["normbud"])["label"]


def _company_render_profile(company_key: Optional[str]) -> dict:
    key = _normalize_company_key(company_key)
    return COMPANY_RENDER_PROFILES.get(key, COMPANY_RENDER_PROFILES["normbud"])


@lru_cache(maxsize=16)
def _asset_data_uri(file_name: str) -> str:
    if not file_name:
        return ""
    path = PREVIEW_ASSETS_DIR / file_name
    if not path.exists():
        return ""
    mime = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".svg": "image/svg+xml",
    }.get(path.suffix.lower(), "application/octet-stream")
    encoded = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def _is_admin(request: Request) -> bool:
    return bool(getattr(request.state, "is_superuser", False))


def _require_admin(request: Request):
    if not _is_admin(request):
        raise HTTPException(status_code=403, detail="Admin access required")


def _id_conditions(column, value):
    variants = []
    try:
        parsed = value if isinstance(value, uuid.UUID) else uuid.UUID(str(value))
        variants.extend([parsed, str(parsed), parsed.hex])
    except (ValueError, TypeError):
        variants.append(str(value))
    return or_(*[column == v for v in variants])


def _parse_date(value: Optional[str]) -> date:
    if not value:
        return date.today()
    if isinstance(value, date):
        return value
    try:
        return datetime.fromisoformat(value).date()
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="Invalid letter_date format, expected YYYY-MM-DD")


def _parse_filter_date(value: Optional[str], label: str) -> Optional[date]:
    if not value:
        return None
    if isinstance(value, date):
        return value
    try:
        return datetime.fromisoformat(value).date()
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail=f"Invalid {label} format, expected YYYY-MM-DD")


def _normalize_company_key(value: Optional[str]) -> str:
    if not value:
        return "normbud"
    raw = str(value).strip().lower()
    mapping = {
        "normbud": "normbud",
        "\u043d\u043e\u0440\u043c\u0431\u0443\u0434": "normbud",
        "bayer": "bayer",
        "\u0431\u0430\u0439\u0435\u0440": "bayer",
        "morozov": "morozov",
        "\u043c\u043e\u0440\u043e\u0437\u043e\u0432": "morozov",
        "\u0438\u043f \u043c\u043e\u0440\u043e\u0437\u043e\u0432": "morozov",
        "ip morozov": "morozov",
    }
    return mapping.get(raw, "normbud")


def _normalize_document_kind(value: Optional[str]) -> str:
    raw = str(value or "letter").strip().lower()
    aliases = {
        "letter": "letter",
        "mail": "letter",
        "invoice": "invoice",
        "bill": "invoice",
        "счет": "invoice",
        "счёт": "invoice",
        "upd": "upd",
        "упд": "upd",
        "act": "act",
        "акт": "act",
        "vat_invoice": "vat_invoice",
        "sf": "vat_invoice",
        "счет-фактура": "vat_invoice",
        "счёт-фактура": "vat_invoice",
    }
    normalized = aliases.get(raw)
    if not normalized:
        raise HTTPException(status_code=400, detail="Invalid document_kind")
    return normalized


def _json_loads_or_default(value, default):
    if value is None or value == "":
        return default
    if isinstance(value, (list, dict)):
        return value
    try:
        return json.loads(value)
    except (TypeError, ValueError):
        return default


def _json_dumps_or_none(value) -> Optional[str]:
    if value is None or value == "":
        return None
    if isinstance(value, str):
        parsed = _json_loads_or_default(value, None)
        if parsed is None:
            return value
        value = parsed
    return json.dumps(value, ensure_ascii=False)


def _normalize_editor_payload(
    *,
    document_kind: str,
    editor_mode: Optional[str] = None,
    editor_schema_version: Optional[int] = None,
    editor_draft=None,
    editor_validation=None,
    editor_render_context=None,
) -> dict:
    normalized_draft = normalize_editor_draft(editor_draft, document_kind)
    validation_payload = editor_validation if isinstance(editor_validation, dict) else validate_editor_draft(normalized_draft, document_kind)
    render_context_payload = editor_render_context if isinstance(editor_render_context, dict) else {"snapshot_mode": "live"}
    return {
        "editor_mode": normalize_editor_mode(editor_mode),
        "editor_schema_version": int(editor_schema_version or normalized_draft.get("schema_version") or 1),
        "editor_draft_json": _json_dumps_or_none(normalized_draft),
        "editor_validation_json": _json_dumps_or_none(validation_payload),
        "editor_render_context_json": _json_dumps_or_none(render_context_payload),
    }


def _resolve_path_value(payload, path: str):
    current = payload
    for part in str(path or "").split("."):
        if not part:
            continue
        if isinstance(current, dict):
            current = current.get(part)
        else:
            return None
    return current


def _render_editor_text_template(template: Optional[str], payload: dict) -> str:
    raw_template = str(template or "").strip()
    if not raw_template:
        return ""

    def replace(match):
        key = str(match.group(1) or "").strip()
        value = _resolve_path_value(payload, key)
        if value is None:
            return ""
        if isinstance(value, (list, dict)):
            return ""
        return str(value)

    return re.sub(r"{{\s*([^}]+?)\s*}}", replace, raw_template).strip()


STRUCTURED_TEMPLATE_BLOCK_DEFS = [
    {
        "key": "intro_paragraph",
        "label": "Вводный абзац",
        "document_kinds": ["letter"],
    },
    {
        "key": "document_basis",
        "label": "Основание документа",
        "document_kinds": ["invoice", "upd", "act", "vat_invoice"],
    },
    {
        "key": "payment_due_date",
        "label": "Срок оплаты",
        "document_kinds": ["invoice"],
    },
    {
        "key": "invoice_items",
        "label": "Табличная часть счета",
        "document_kinds": ["invoice"],
    },
    {
        "key": "stage_lines",
        "label": "Строки по этапам",
        "document_kinds": ["upd", "act", "vat_invoice"],
    },
    {
        "key": "payment_allocations",
        "label": "Зачеты платежей",
        "document_kinds": ["act"],
    },
    {
        "key": "totals",
        "label": "Итоги документа",
        "document_kinds": ["invoice", "upd", "act", "vat_invoice"],
    },
    {
        "key": "signature_stamp",
        "label": "Подписи и печать",
        "document_kinds": ["letter", "invoice", "upd", "act", "vat_invoice"],
    },
]


def _get_structured_template_blocks(document_kind: str) -> list[dict]:
    normalized_kind = _normalize_document_kind(document_kind)
    return [
        item
        for item in STRUCTURED_TEMPLATE_BLOCK_DEFS
        if normalized_kind in item.get("document_kinds", [])
    ]


def _replace_editor_field_tokens(raw_html: str, payload: dict) -> str:
    if not raw_html:
        return ""

    def replace_chip(match):
        tag_html = str(match.group(0) or "")
        key_match = re.search(r"data-field-key=(?:\"([^\"]+)\"|'([^']+)')", tag_html, flags=re.IGNORECASE)
        key = html.unescape(str((key_match.group(1) if key_match else "") or (key_match.group(2) if key_match else "") or "").strip())
        value = _resolve_path_value(payload, key)
        if value is None:
            return ""
        if isinstance(value, (list, dict)):
            return ""
        return html.escape(str(value))

    html_value = re.sub(
        r"<span\b[^>]*data-field-chip=(?:\"true\"|'true')[^>]*>.*?</span>",
        replace_chip,
        raw_html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return _render_editor_text_template(html_value, payload)


def _build_structured_stages_html(payload: dict) -> str:
    stages = payload.get("stages") or []
    if not isinstance(stages, list) or not stages:
        return ""
    rows = []
    for index, stage in enumerate(stages, start=1):
        if not isinstance(stage, dict):
            continue
        rows.append(
            "<tr>"
            f"<td>{index}</td>"
            f"<td>{html.escape(str(stage.get('name') or ''))}</td>"
            f"<td>{html.escape(str(stage.get('amount') or ''))}</td>"
            "</tr>"
        )
    if not rows:
        return ""
    return (
        "<div class=\"structured-fragment structured-fragment--stages\">"
        "<p><strong>Этапы</strong></p>"
        "<table><thead><tr><th>№</th><th>Этап</th><th>Сумма</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody></table>"
        "</div>"
    )


def _build_structured_payment_allocations_html(payload: dict) -> str:
    payments = payload.get("linked_payment_items") or []
    if not isinstance(payments, list) or not payments:
        return ""
    rows = []
    for index, item in enumerate(payments, start=1):
        if not isinstance(item, dict):
            continue
        rows.append(
            "<tr>"
            f"<td>{index}</td>"
            f"<td>{html.escape(str(item.get('entry_id') or ''))}</td>"
            f"<td>{html.escape(str(item.get('amount') or ''))}</td>"
            f"<td>{html.escape(str(item.get('note') or ''))}</td>"
            "</tr>"
        )
    if not rows:
        return ""
    return (
        "<div class=\"structured-fragment structured-fragment--payments\">"
        "<p><strong>Зачеты платежей</strong></p>"
        "<table><thead><tr><th>№</th><th>Платеж</th><th>Сумма</th><th>Комментарий</th></tr></thead>"
        f"<tbody>{''.join(rows)}</tbody></table>"
        "</div>"
    )


def _build_structured_totals_html(payload: dict) -> str:
    document_payload = payload.get("document") or {}
    total_amount = html.escape(str(document_payload.get("total_amount") or ""))
    total_words = html.escape(str(document_payload.get("total_amount_words") or ""))
    vat_amount = html.escape(str(document_payload.get("vat_amount") or ""))
    vat_rate = html.escape(str(document_payload.get("vat_rate_text") or document_payload.get("vat_rate") or ""))
    if not any([total_amount, total_words, vat_amount, vat_rate]):
        return ""
    return (
        "<div class=\"structured-fragment structured-fragment--totals\">"
        "<p><strong>Итоги документа</strong></p>"
        "<table><tbody>"
        f"<tr><th>Сумма без НДС</th><td>{html.escape(str(document_payload.get('amount_without_vat') or ''))}</td></tr>"
        f"<tr><th>НДС ({vat_rate}%)</th><td>{vat_amount}</td></tr>"
        f"<tr><th>Итого</th><td>{total_amount}</td></tr>"
        f"<tr><th>Сумма прописью</th><td>{total_words}</td></tr>"
        "</tbody></table>"
        "</div>"
    )


def _build_structured_signature_html(payload: dict) -> str:
    our_company = payload.get("our_company") or {}
    signer_title = html.escape(str(our_company.get("signer_title") or payload.get("signer_title") or "Руководитель"))
    signer_name = html.escape(str(our_company.get("signer_name") or payload.get("signer_name") or ""))
    if not signer_title and not signer_name:
        return ""
    return (
        "<div class=\"structured-fragment structured-fragment--signature\">"
        "<table><tbody>"
        f"<tr><th>{signer_title}</th><td>{signer_name}</td></tr>"
        "</tbody></table>"
        "</div>"
    )


def _build_structured_invoice_items_html(payload: dict) -> str:
    document_payload = payload.get("document") or {}
    basis = html.escape(str(document_payload.get("basis") or payload.get("subject") or ""))
    amount = html.escape(str(document_payload.get("total_amount") or ""))
    vat_rate = html.escape(str(document_payload.get("vat_rate_text") or document_payload.get("vat_rate") or ""))
    if not any([basis, amount]):
        return ""
    return (
        "<div class=\"structured-fragment structured-fragment--invoice\">"
        "<table><thead><tr><th>№</th><th>Наименование</th><th>Ставка НДС</th><th>Сумма</th></tr></thead>"
        "<tbody>"
        f"<tr><td>1</td><td>{basis or 'Оплата по документу'}</td><td>{vat_rate}</td><td>{amount}</td></tr>"
        "</tbody></table>"
        "</div>"
    )


def _build_structured_template_block_html(block_key: str, payload: dict) -> str:
    normalized_key = str(block_key or "").strip().lower()
    document_payload = payload.get("document") or {}
    if normalized_key == "intro_paragraph":
        intro = _build_auto_intro_paragraph(payload)
        return _plain_text_to_html(intro)
    if normalized_key == "document_basis":
        basis = str(document_payload.get("basis") or "").strip()
        return f"<p><strong>Основание:</strong> {html.escape(basis)}</p>" if basis else ""
    if normalized_key == "payment_due_date":
        payment_due_date = str(document_payload.get("payment_due_date") or "").strip()
        return f"<p><strong>Оплатить не позднее:</strong> {html.escape(payment_due_date)}</p>" if payment_due_date else ""
    if normalized_key == "invoice_items":
        return _build_structured_invoice_items_html(payload)
    if normalized_key == "stage_lines":
        return _build_structured_stages_html(payload)
    if normalized_key == "payment_allocations":
        return _build_structured_payment_allocations_html(payload)
    if normalized_key == "totals":
        return _build_structured_totals_html(payload)
    if normalized_key == "signature_stamp":
        return _build_structured_signature_html(payload)
    return ""


def _replace_editor_template_blocks(raw_html: str, payload: dict) -> str:
    if not raw_html:
        return ""

    def replace_template_block(match):
        tag_html = str(match.group(0) or "")
        key_match = re.search(r"data-template-key=(?:\"([^\"]+)\"|'([^']+)')", tag_html, flags=re.IGNORECASE)
        key = html.unescape(str((key_match.group(1) if key_match else "") or (key_match.group(2) if key_match else "") or "").strip())
        return _build_structured_template_block_html(key, payload)

    return re.sub(
        r"<div\b[^>]*data-template-block=(?:\"true\"|'true')[^>]*>.*?</div>",
        replace_template_block,
        raw_html,
        flags=re.IGNORECASE | re.DOTALL,
    )


def _render_editor_rich_text_html(raw_html: Optional[str], payload: dict) -> str:
    html_value = str(raw_html or "").strip()
    if not html_value:
        return ""
    html_value = _replace_editor_template_blocks(html_value, payload)
    html_value = _replace_editor_field_tokens(html_value, payload)
    # H1: user-authored rich-text must pass the allowlist sanitizer before it
    # reaches any render path (xhtml2pdf / soffice / node) — otherwise
    # `<img src="file:///...">` / remote URLs cause local-file read & SSRF.
    return _sanitize_html(html_value.strip())


def _build_structured_render_sections(payload: dict, editor_draft: dict) -> list[dict]:
    document_kind = str(payload.get("document", {}).get("kind") or "letter")
    sections: list[dict] = []
    for block in get_editor_blocks(editor_draft, document_kind):
        block_type = str(block.get("type") or "")
        attrs = block.get("attrs") if isinstance(block.get("attrs"), dict) else {}
        html_fragment = ""
        if block_type == "intro_paragraph":
            text = ""
            if str(attrs.get("mode") or "").strip() == "contract_intro" and not str(attrs.get("text") or "").strip():
                text = _build_auto_intro_paragraph(payload)
            else:
                text = _render_editor_text_template(attrs.get("text"), payload)
            html_fragment = _plain_text_to_html(text)
        elif block_type == "basis_block":
            basis = str(payload.get("document", {}).get("basis") or "").strip()
            if basis:
                html_fragment = f"<p><strong>Основание:</strong> {html.escape(basis)}</p>"
        elif block_type == "rich_text_block":
            html_fragment = _render_editor_rich_text_html(attrs.get("html"), payload)
        elif block_type == "invoice_items_table":
            html_fragment = _build_structured_invoice_items_html(payload)
        elif block_type == "stage_lines_block":
            html_fragment = _build_structured_stages_html(payload)
        elif block_type == "payment_allocation_block":
            html_fragment = _build_structured_payment_allocations_html(payload)
        elif block_type == "totals_block":
            html_fragment = _build_structured_totals_html(payload)
        elif block_type == "signature_stamp":
            html_fragment = _build_structured_signature_html(payload)
        if html_fragment:
            sections.append(
                {
                    "id": str(block.get("id") or ""),
                    "type": block_type,
                    "html": html_fragment,
                }
            )
    return sections


def _plain_text_to_html(text: Optional[str]) -> str:
    value = str(text or "").strip()
    if not value:
        return ""
    paragraphs = []
    for raw_paragraph in re.split(r"\n\s*\n", value):
        paragraph = raw_paragraph.strip()
        if not paragraph:
            continue
        paragraphs.append(f"<p>{html.escape(paragraph).replace(chr(10), '<br>')}</p>")
    return "".join(paragraphs)


def _build_auto_intro_paragraph(payload: dict) -> str:
    our_company_name = (
        _first_value(
            _resolve_path_value(payload, "our_company.full_name"),
            _resolve_path_value(payload, "our_company.name"),
        )
        or "Нашей компанией"
    )
    recipient_name = (
        _first_value(
            _resolve_path_value(payload, "recipient.short_name"),
            _resolve_path_value(payload, "recipient.name"),
        )
        or "получателем"
    )
    contract_number = _first_value(_resolve_path_value(payload, "contract.number"))
    contract_date = _first_value(_resolve_path_value(payload, "contract.date"))
    deal_object = _first_value(
        _resolve_path_value(payload, "deal.obj_name"),
        _resolve_path_value(payload, "deal.title"),
    )

    number_part = _contract_number_with_prefix(contract_number) if contract_number else ""
    date_part = f" от {contract_date}" if contract_date else ""
    object_part = f' по объекту «{deal_object}»' if deal_object else ""
    if number_part or date_part:
        return (
            f"Между {our_company_name} и {recipient_name} заключен Договор "
            f"{number_part}{date_part}{object_part}."
        ).strip()
    if object_part:
        return f"Между {our_company_name} и {recipient_name} ведутся работы{object_part}.".strip()
    return f"Между {our_company_name} и {recipient_name} ведется взаимодействие в рамках текущего проекта."


def _resolve_payment_due_date(letter_date: Optional[date], attrs: dict) -> Optional[date]:
    if not letter_date:
        return None
    mode = str(attrs.get("payment_due_date_mode") or "workdays_plus").strip().lower()
    if mode == "document_date":
        return letter_date
    if mode == "workdays_plus":
        try:
            workdays = int(attrs.get("payment_due_days") or 0)
        except (TypeError, ValueError):
            workdays = 0
        return _add_workdays(letter_date, max(workdays, 0))
    manual_value = attrs.get("payment_due_date")
    if manual_value:
        try:
            return datetime.fromisoformat(str(manual_value)).date()
        except (TypeError, ValueError):
            return None
    return None


def _collect_structured_body_html(payload: dict, editor_draft: dict) -> str:
    sections = _build_structured_render_sections(payload, editor_draft)
    return "\n".join(
        str(item.get("html") or "").strip()
        for item in sections
        if str(item.get("html") or "").strip()
    ).strip()


def _apply_structured_editor_overrides(document: OutgoingDocument, payload: dict, editor_draft: dict) -> dict:
    next_payload = json.loads(json.dumps(payload, ensure_ascii=False))
    document_kind = str(next_payload.get("document", {}).get("kind") or _normalize_document_kind(getattr(document, "document_kind", None)))

    meta_block = get_first_editor_block(editor_draft, document_kind, "document_meta")
    if meta_block:
        attrs = meta_block.get("attrs") if isinstance(meta_block.get("attrs"), dict) else {}
        payment_due_date = _resolve_payment_due_date(document.letter_date, attrs)
        if payment_due_date:
            next_payload["document"]["payment_due_date"] = _format_ru_date(payment_due_date)
            next_payload["document"]["payment_due_date_iso"] = payment_due_date.isoformat()

    basis_block = get_first_editor_block(editor_draft, document_kind, "basis_block")
    if basis_block:
        attrs = basis_block.get("attrs") if isinstance(basis_block.get("attrs"), dict) else {}
        mode = str(attrs.get("mode") or "contract_auto").strip().lower()
        if mode in {"manual", "template"}:
            rendered_basis = _render_editor_text_template(attrs.get("text_pattern"), next_payload)
            if rendered_basis:
                next_payload["document"]["basis"] = rendered_basis

    structured_sections = _build_structured_render_sections(next_payload, editor_draft)
    structured_body = "\n".join(
        str(item.get("html") or "").strip()
        for item in structured_sections
        if str(item.get("html") or "").strip()
    ).strip()
    next_payload["body"] = structured_body or ""

    next_payload["subject"] = next_payload.get("document", {}).get("subject") or next_payload.get("subject") or ""
    next_payload["editor"] = {
        "mode": "structured",
        "schema_version": int(editor_draft.get("schema_version") or 1),
        "draft": editor_draft,
        "block_types": [block.get("type") for block in get_editor_blocks(editor_draft, document_kind)],
        "render": {
            "sections": structured_sections,
            "body_html": structured_body or "",
        },
        "resolved": {
            "basis": next_payload.get("document", {}).get("basis") or "",
            "payment_due_date": next_payload.get("document", {}).get("payment_due_date") or "",
            "body_html": structured_body or "",
        },
    }
    return next_payload


def _build_editor_resolved_fields(render_payload: dict, document_kind: str) -> dict:
    resolved = {}
    for item in get_template_fields(module="outgoing_registry", document_kind=document_kind):
        key = str(item.get("key") or "")
        if not key:
            continue
        value = _resolve_path_value(render_payload, key)
        if value is None:
            if key == "stages":
                value = f"[{len(render_payload.get('stages') or [])} items]"
            elif key.startswith("stage."):
                first_stage = (render_payload.get("stages") or [{}])[0] if isinstance(render_payload.get("stages"), list) else {}
                value = first_stage.get(key.split(".", 1)[1], "")
            elif key == "linked_payment_items":
                value = f"[{len(render_payload.get('linked_payment_items') or [])} items]"
            elif key.startswith("payment."):
                first_payment = (render_payload.get("linked_payment_items") or [{}])[0] if isinstance(render_payload.get("linked_payment_items"), list) else {}
                value = first_payment.get(key.split(".", 1)[1], "")
            else:
                value = ""
        resolved[key] = value
    return resolved


def _normalize_compare_text(value: Optional[str]) -> str:
    if not value:
        return ""
    normalized = unicodedata.normalize("NFKD", str(value)).casefold()
    normalized = normalized.replace("«", " ").replace("»", " ").replace('"', " ")
    normalized = re.sub(r"[().,;:_\-]+", " ", normalized)
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized.strip()


def _contract_number_with_prefix(value: Optional[str]) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    return raw if raw.startswith("№") else f"№{raw}"


def _format_ru_date(value: Optional[date]) -> str:
    if not value:
        return ""
    if isinstance(value, datetime):
        value = value.date()
    return value.strftime("%d.%m.%Y")


def _format_money(value) -> str:
    try:
        amount = round(float(value or 0), 2)
    except (TypeError, ValueError):
        amount = 0.0
    return f"{amount:,.2f}".replace(",", " ").replace(".", ",")


def _format_rate(value) -> str:
    try:
        amount = round(float(value or 0), 2)
    except (TypeError, ValueError):
        amount = 0.0
    if amount.is_integer():
        return str(int(amount))
    return str(amount).replace(".", ",")


def _plural_ru(value: int, one: str, few: str, many: str) -> str:
    value = abs(int(value))
    if 11 <= value % 100 <= 14:
        return many
    tail = value % 10
    if tail == 1:
        return one
    if 2 <= tail <= 4:
        return few
    return many


def _number_to_words_ru(value: int) -> str:
    value = int(value or 0)
    if value == 0:
        return "ноль"

    units_m = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    units_f = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать", "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста", "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    scales = [
        ("", "", "", "m"),
        ("тысяча", "тысячи", "тысяч", "f"),
        ("миллион", "миллиона", "миллионов", "m"),
        ("миллиард", "миллиарда", "миллиардов", "m"),
    ]

    def triad_words(number: int, gender: str) -> List[str]:
        result = []
        h = number // 100
        t = (number % 100) // 10
        u = number % 10
        if h:
            result.append(hundreds[h])
        if t == 1:
            result.append(teens[u])
        else:
            if t:
                result.append(tens[t])
            if u:
                result.append((units_f if gender == "f" else units_m)[u])
        return result

    parts = []
    group_index = 0
    while value > 0:
        group_value = value % 1000
        if group_value:
            scale_one, scale_few, scale_many, gender = scales[group_index]
            words = triad_words(group_value, gender)
            if group_index:
                words.append(_plural_ru(group_value, scale_one, scale_few, scale_many))
            parts = words + parts
        value //= 1000
        group_index += 1
    return " ".join(parts)


def _amount_to_words_ru(value) -> str:
    try:
        amount = round(float(value or 0), 2)
    except (TypeError, ValueError):
        amount = 0.0
    rubles = int(amount)
    kopeks = int(round((amount - rubles) * 100))
    if kopeks == 100:
        rubles += 1
        kopeks = 0
    words = _number_to_words_ru(rubles).capitalize()
    ruble_label = _plural_ru(rubles, "рубль", "рубля", "рублей")
    kopek_label = _plural_ru(kopeks, "копейка", "копейки", "копеек")
    return f"{words} {ruble_label} {kopeks:02d} {kopek_label}"


def _add_workdays(start_date: Optional[date], days: int) -> Optional[date]:
    if not start_date:
        return None
    current = start_date.date() if isinstance(start_date, datetime) else start_date
    added = 0
    while added < days:
        current += timedelta(days=1)
        if current.weekday() < 5:
            added += 1
    return current


def _safe_float(value) -> float:
    try:
        return float(value or 0)
    except (TypeError, ValueError):
        return 0.0


def _company_matches_key(company: Optional[Company], company_key: Optional[str]) -> bool:
    if not company:
        return False
    key = _normalize_company_key(company_key)
    aliases = FORMAL_COMPANY_ALIASES.get(key, [])
    haystack = " ".join(
        _normalize_compare_text(part)
        for part in (company.short_name, company.name, company.full_name)
        if part
    )
    return any(alias in haystack for alias in aliases)


def _first_value(*values) -> str:
    for value in values:
        if value is None:
            continue
        normalized = str(value).strip()
        if normalized:
            return normalized
    return ""


def _director_short_name(value: Optional[str]) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    raw = re.sub(r"\s+", " ", raw).strip()
    initials_first = re.match(r"^([А-ЯЁA-Z])\.\s*([А-ЯЁA-Z])\.?\s+(.+)$", raw, re.IGNORECASE)
    if initials_first:
        first, second, surname = initials_first.groups()
        return f"{surname.strip()} {first.upper()}. {second.upper()}."
    one_initial_first = re.match(r"^([А-ЯЁA-Z])\.?\s+(.+)$", raw, re.IGNORECASE)
    if one_initial_first:
        first, surname = one_initial_first.groups()
        return f"{surname.strip()} {first.upper()}."
    parts = [part for part in raw.split(" ") if part]
    if len(parts) >= 3:
        surname, name, patronymic = parts[0], parts[1], parts[2]
        return f"{surname} {name[:1]}. {patronymic[:1]}.".strip()
    return raw


def _bank_account_payload(account: Optional[dict]) -> dict:
    account = account if isinstance(account, dict) else {}
    name = _first_value(account.get("bank_name"), account.get("bank"), account.get("name"), account.get("bankName"))
    bik = _first_value(account.get("bik"), account.get("bic"))
    rs = _first_value(account.get("account_rs"), account.get("rs"), account.get("account"), account.get("settlement_account"))
    ks = _first_value(account.get("account_ks"), account.get("ks"), account.get("correspondent_account"))
    return {
        "name": name,
        "bank_name": name,
        "bik": bik,
        "rs": rs,
        "account_rs": rs,
        "ks": ks,
        "account_ks": ks,
        "correspondent_account": ks,
    }


def _company_payload(
    company: Optional[Company],
    *,
    fallback: Optional[dict] = None,
    bank_account: Optional[dict] = None,
    director_name: str = "",
) -> dict:
    fallback = fallback or {}
    bank = _bank_account_payload(bank_account)
    return {
        "name": _first_value(getattr(company, "name", None), fallback.get("name")),
        "short_name": _first_value(getattr(company, "short_name", None), fallback.get("short_name"), getattr(company, "name", None), fallback.get("name")),
        "full_name": _first_value(getattr(company, "full_name", None), fallback.get("full_name"), getattr(company, "name", None), fallback.get("name")),
        "inn": _first_value(getattr(company, "inn", None), fallback.get("inn")),
        "kpp": _first_value(getattr(company, "kpp", None), fallback.get("kpp")),
        "address": _first_value(getattr(company, "address", None), fallback.get("address")),
        "director_name": director_name,
        "bank": bank,
    }


async def _load_our_company_for_document(
    db: AsyncSession,
    company_key: Optional[str],
    deal: Optional[Deal],
) -> Optional[Company]:
    key = _normalize_company_key(company_key)
    fallback = OUR_COMPANY_REQUISITES.get(key, {})
    company = await Company.get_by_inn(db, fallback.get("inn")) if fallback.get("inn") else None
    if company:
        return company
    if deal and getattr(deal, "our_company_id", None):
        candidate = await Company.get_by_id(db, str(deal.our_company_id))
        if candidate and (_company_matches_key(candidate, key) or candidate.inn == fallback.get("inn")):
            return candidate
    return None


def _display_outgoing_number(value: Optional[str]) -> str:
    if not value:
        return ""
    if ":" in value:
        return value.rsplit(":", 1)[1]
    return value


async def _get_company_sequence(db: AsyncSession, company_key: str) -> int:
    key = _normalize_company_key(company_key)
    for _ in range(3):
        result = await db.execute(
            select(OutgoingNumberSequence)
            .where(OutgoingNumberSequence.our_company_key == key)
            .with_for_update()
        )
        row = result.scalar_one_or_none()
        if not row:
            row = OutgoingNumberSequence(our_company_key=key, next_seq=settings.OUTGOING_NUMBER_START)
            db.add(row)
            try:
                await db.commit()
            except IntegrityError:
                await db.rollback()
                continue
        seq = int(row.next_seq)
        row.next_seq = seq + 1
        await db.commit()
        return seq
    raise HTTPException(status_code=409, detail="Unable to allocate company sequence, retry")

def _escape_text(value: Optional[str]) -> str:
    if not value:
        return ""
    return html.escape(value, quote=True)


ALLOWED_HTML_TAGS = {
    "b", "strong", "i", "em", "u", "br", "p", "ul", "ol", "li",
    "div", "span", "table", "thead", "tbody", "tr", "th", "td",
}
ALLOWED_HTML_ATTRS = {
    "table": {"style"},
    "th": {"style", "colspan", "rowspan"},
    "td": {"style", "colspan", "rowspan"},
}
ALLOWED_STYLE_PROPS = {
    "width",
    "border",
    "border-top",
    "border-right",
    "border-bottom",
    "border-left",
    "border-collapse",
    "padding",
    "margin",
    "margin-top",
    "margin-right",
    "margin-bottom",
    "margin-left",
    "background",
    "background-color",
    "vertical-align",
    "font-weight",
    "text-align",
    "table-layout",
    "word-wrap",
}


def _sanitize_style(style_value: Optional[str]) -> str:
    if not style_value:
        return ""
    safe_parts: List[str] = []
    for chunk in style_value.split(";"):
        if ":" not in chunk:
            continue
        prop, raw_value = chunk.split(":", 1)
        prop = prop.strip().lower()
        value = raw_value.strip()
        if not prop or not value or prop not in ALLOWED_STYLE_PROPS:
            continue
        safe_parts.append(f"{prop}: {value}")
    return "; ".join(safe_parts)


class _SafeHtmlParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self._parts: List[str] = []

    def handle_starttag(self, tag, attrs):
        if tag in ALLOWED_HTML_TAGS:
            if tag == "br":
                self._parts.append("<br>")
                return
            safe_attrs: List[str] = []
            allowed_attrs = ALLOWED_HTML_ATTRS.get(tag, set())
            for attr_name, attr_value in attrs:
                if attr_name not in allowed_attrs:
                    continue
                if attr_name == "style":
                    sanitized_style = _sanitize_style(attr_value)
                    if sanitized_style:
                        safe_attrs.append(f'style="{html.escape(sanitized_style, quote=True)}"')
                    continue
                if attr_name in {"colspan", "rowspan"} and attr_value and str(attr_value).isdigit():
                    safe_attrs.append(f'{attr_name}="{attr_value}"')
            attrs_html = f" {' '.join(safe_attrs)}" if safe_attrs else ""
            self._parts.append(f"<{tag}{attrs_html}>")

    def handle_endtag(self, tag):
        if tag in ALLOWED_HTML_TAGS and tag != "br":
            self._parts.append(f"</{tag}>")

    def handle_data(self, data):
        self._parts.append(_escape_text(data))

    def get_html(self) -> str:
        return "".join(self._parts)


def _sanitize_html(value: Optional[str]) -> str:
    if not value:
        return ""
    parser = _SafeHtmlParser()
    parser.feed(value)
    parser.close()
    return parser.get_html()


async def _get_next_number(db: AsyncSession) -> int:
    result = await db.execute(select(func.max(OutgoingDocument.outgoing_number_seq)))
    max_value = result.scalar()
    if max_value is None:
        return settings.OUTGOING_NUMBER_START
    return int(max_value) + 1


def _format_outgoing_number(seq: int, letter_date: date) -> str:
    return f"{seq}/{letter_date:%Y-%m}"


def _format_financial_number(seq: int, document_date: date) -> str:
    return f"{document_date:%d%m%y}/{int(seq):05d}"


async def _get_daily_financial_sequence(
    db: AsyncSession,
    *,
    company_key: str,
    document_kind: str,
    document_date: date,
) -> int:
    company_key = _normalize_company_key(company_key)
    document_kind = _normalize_document_kind(document_kind)
    for _ in range(3):
        result = await db.execute(
            select(OutgoingDailyNumberSequence)
            .where(
                and_(
                    OutgoingDailyNumberSequence.our_company_key == company_key,
                    OutgoingDailyNumberSequence.document_kind == document_kind,
                    OutgoingDailyNumberSequence.sequence_date == document_date,
                )
            )
            .with_for_update()
        )
        row = result.scalar_one_or_none()
        if not row:
            max_existing = await _get_max_daily_financial_sequence(
                db,
                company_key=company_key,
                document_kind=document_kind,
                document_date=document_date,
            )
            row = OutgoingDailyNumberSequence(
                id=str(uuid.uuid4()),
                our_company_key=company_key,
                document_kind=document_kind,
                sequence_date=document_date,
                next_seq=max_existing + 1,
            )
            db.add(row)
            try:
                await db.commit()
            except IntegrityError:
                await db.rollback()
                continue
        seq = int(row.next_seq)
        row.next_seq = seq + 1
        await db.commit()
        return seq
    raise HTTPException(status_code=409, detail="Unable to allocate daily financial sequence, retry")


async def _get_max_daily_financial_sequence(
    db: AsyncSession,
    *,
    company_key: str,
    document_kind: str,
    document_date: date,
) -> int:
    result = await db.execute(
        select(func.max(OutgoingDocument.outgoing_number_company_seq)).where(
            and_(
                OutgoingDocument.our_company_key == company_key,
                OutgoingDocument.document_kind == document_kind,
                OutgoingDocument.letter_date == document_date,
            )
        )
    )
    max_value = result.scalar()
    return int(max_value or 0)


async def _peek_daily_financial_sequence(
    db: AsyncSession,
    *,
    company_key: str,
    document_kind: str,
    document_date: date,
) -> int:
    company_key = _normalize_company_key(company_key)
    document_kind = _normalize_document_kind(document_kind)
    result = await db.execute(
        select(OutgoingDailyNumberSequence).where(
            and_(
                OutgoingDailyNumberSequence.our_company_key == company_key,
                OutgoingDailyNumberSequence.document_kind == document_kind,
                OutgoingDailyNumberSequence.sequence_date == document_date,
            )
        )
    )
    row = result.scalar_one_or_none()
    if row:
        return int(row.next_seq)
    max_existing = await _get_max_daily_financial_sequence(
        db,
        company_key=company_key,
        document_kind=document_kind,
        document_date=document_date,
    )
    return max_existing + 1


async def _get_next_contract_act_number(db: AsyncSession, contract_id: str) -> int:
    documents = await ContractDocument.get_by_contract_and_type(db, contract_id, "act")
    return max((int(doc.number_in_contract or 0) for doc in documents), default=0) + 1

def _format_date_long(value: date) -> str:
    months = [
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
    ]
    return f"{value.day:02d} {months[value.month - 1]} {value.year}"

def _parse_outgoing_suffix(value: Optional[str]) -> Optional[str]:
    if not value:
        return None
    suffix = str(value).strip()
    if not re.match(r"^\d{4}-\d{2}$", suffix):
        raise HTTPException(status_code=422, detail="Invalid outgoing number suffix format, expected YYYY-MM")
    return suffix

def _build_outgoing_number_with_suffix(document: OutgoingDocument, suffix: str) -> str:
    display = _display_outgoing_number(document.outgoing_number) or ""
    match = re.match(r"^(\d+)[/\\]", display)
    if not match:
        raise HTTPException(status_code=422, detail="Unable to parse outgoing number prefix")
    seq_part = match.group(1)
    new_display = f"{seq_part}/{suffix}"
    company_key = document.our_company_key or "normbud"
    if ":" in document.outgoing_number:
        return f"{company_key}:{new_display}"
    return new_display

def _format_outgoing_number_for_file(value: Optional[str]) -> str:
    if not value:
        return "draft"
    match = re.match(r"^(\d+)[/\\](\d{4})-(\d{2})$", value)
    if match:
        return f"{match.group(1)}_{match.group(3)}-{match.group(2)}"
    return re.sub(r"[/\\\\]", "_", value)


def _build_outgoing_file_base(document: OutgoingDocument) -> str:
    label = _company_label(document.our_company_key).replace(" ", "_")
    display_number = _display_outgoing_number(document.outgoing_number) or document.outgoing_number
    number_part = _format_outgoing_number_for_file(display_number)
    kind_label = DOCUMENT_KIND_LABELS.get(_normalize_document_kind(getattr(document, "document_kind", None)), "Документ")
    return f"{label}_{kind_label}_{number_part}"


def _build_paths(outgoing_number: str) -> dict:
    base = (settings.STORAGE_LOCAL_ROOT or "").rstrip("/")
    if ":" in outgoing_number:
        safe_id = clean_name(outgoing_number.replace(":", "_"))
    else:
        safe_id = clean_name(outgoing_number)
    root = f"{base}/Outgoing/{safe_id}"
    return {
        "root": root,
        "attachments": f"{root}/Attachments",
        "versions": f"{root}/Versions",
    }


def _build_current_render_paths(document: OutgoingDocument) -> dict:
    paths = _build_paths(document.outgoing_number)
    base_name = _build_outgoing_file_base_clean(document)
    current_root = f"{paths['root']}/Current"
    return {
        "root": current_root,
        "docx_name": clean_name(f"{base_name}_current.docx"),
        "pdf_name": clean_name(f"{base_name}_current.pdf"),
        "docx_path": f"{current_root}/{clean_name(f'{base_name}_current.docx')}",
        "pdf_path": f"{current_root}/{clean_name(f'{base_name}_current.pdf')}",
    }


async def _get_current_render_files(db: AsyncSession, document_id: str) -> List[OutgoingDocumentFile]:
    result = await db.execute(
        select(OutgoingDocumentFile)
        .where(
            and_(
                OutgoingDocumentFile.document_id == str(document_id),
                OutgoingDocumentFile.version_id.is_(None),
                OutgoingDocumentFile.file_type.in_(("render_docx", "render_pdf")),
            )
        )
        .order_by(OutgoingDocumentFile.created_at.desc())
    )
    return result.scalars().all()


async def _get_current_render_file(
    db: AsyncSession,
    document_id: str,
    file_type: str,
) -> Optional[OutgoingDocumentFile]:
    result = await db.execute(
        select(OutgoingDocumentFile)
        .where(
            and_(
                OutgoingDocumentFile.document_id == str(document_id),
                OutgoingDocumentFile.version_id.is_(None),
                OutgoingDocumentFile.file_type == file_type,
            )
        )
        .order_by(OutgoingDocumentFile.created_at.desc())
    )
    return result.scalars().first()


async def _get_latest_version_bundle(
    db: AsyncSession,
    document_id: str,
):
    versions = await OutgoingDocumentVersion.get_by_document(db, str(document_id))
    if not versions:
        return None, None, None
    latest_version = versions[0]
    result = await db.execute(
        select(OutgoingDocumentFile).where(OutgoingDocumentFile.version_id == str(latest_version.id))
    )
    version_files = result.scalars().all()
    latest_docx = next((item for item in version_files if item.file_type == "docx"), None)
    latest_pdf = next((item for item in version_files if item.file_type == "pdf"), None)
    return latest_version, latest_docx, latest_pdf


def _document_has_changes_after_latest_version(
    document: OutgoingDocument,
    latest_version: Optional[OutgoingDocumentVersion],
) -> bool:
    if not latest_version or not latest_version.created_at:
        return True
    last_document_change = document.updated_at or document.created_at
    if not last_document_change:
        return False
    return last_document_change > latest_version.created_at


def _document_has_changes_after_current_render(
    document: OutgoingDocument,
    current_docx: Optional[OutgoingDocumentFile],
    current_pdf: Optional[OutgoingDocumentFile],
) -> bool:
    last_document_change = document.updated_at or document.created_at
    if not last_document_change:
        return False
    render_timestamps = [
        item.created_at
        for item in (current_docx, current_pdf)
        if item is not None and item.created_at is not None
    ]
    if not render_timestamps:
        return True
    return last_document_change > max(render_timestamps)


def _render_engine_has_changes_after_current_render(
    current_docx: Optional[OutgoingDocumentFile],
    current_pdf: Optional[OutgoingDocumentFile],
) -> bool:
    render_timestamps = [
        item.created_at
        for item in (current_docx, current_pdf)
        if item is not None and item.created_at is not None
    ]
    if not render_timestamps:
        return True
    try:
        engine_changed_at = datetime.fromtimestamp(RENDER_SCRIPT_PATH.stat().st_mtime)
    except OSError:
        return False
    try:
        return engine_changed_at > max(render_timestamps)
    except TypeError:
        return True


async def _read_optional_storage_file(item: Optional[OutgoingDocumentFile]) -> Optional[bytes]:
    if not item or not item.file_path:
        return None
    try:
        return await read_file_bytes(item.file_path)
    except FileNotFoundError:
        return None


async def _resolve_effective_render_payload(
    db: AsyncSession,
    document: OutgoingDocument,
) -> dict:
    current_docx = await _get_current_render_file(db, str(document.id), "render_docx")
    current_pdf = await _get_current_render_file(db, str(document.id), "render_pdf")
    latest_version, latest_docx, latest_pdf = await _get_latest_version_bundle(db, str(document.id))
    active_template_version = await _get_active_document_template_version(db, document)

    current_docx_bytes = await _read_optional_storage_file(current_docx)
    current_pdf_bytes = await _read_optional_storage_file(current_pdf)
    current_render_available = bool(current_docx_bytes or current_pdf_bytes)
    current_render_is_stale = current_render_available and _document_has_changes_after_current_render(
        document,
        current_docx,
        current_pdf,
    )
    if current_render_available:
        current_render_is_stale = current_render_is_stale or _render_engine_has_changes_after_current_render(
            current_docx,
            current_pdf,
        )
    if current_render_available and active_template_version and active_template_version.created_at:
        render_timestamps = [
            item.created_at
            for item in (current_docx, current_pdf)
            if item is not None and item.created_at is not None
        ]
        try:
            current_render_is_stale = current_render_is_stale or (
                bool(render_timestamps) and active_template_version.created_at > max(render_timestamps)
            )
        except TypeError:
            current_render_is_stale = True
    if current_render_available and not current_render_is_stale:
        return {
            "docx_bytes": current_docx_bytes,
            "pdf_bytes": current_pdf_bytes,
            "docx_filename": current_docx.file_name if current_docx else None,
            "pdf_filename": current_pdf.file_name if current_pdf else None,
            "source": "current_render",
        }

    can_fallback_to_latest_version = (
        latest_version is not None
        and not _document_has_changes_after_latest_version(document, latest_version)
    )
    if can_fallback_to_latest_version and active_template_version and active_template_version.created_at:
        try:
            can_fallback_to_latest_version = active_template_version.created_at <= latest_version.created_at
        except TypeError:
            can_fallback_to_latest_version = False
    if can_fallback_to_latest_version:
        latest_docx_bytes = await _read_optional_storage_file(latest_docx)
        latest_pdf_bytes = await _read_optional_storage_file(latest_pdf)
        if latest_docx_bytes or latest_pdf_bytes:
            return {
                "docx_bytes": latest_docx_bytes,
                "pdf_bytes": latest_pdf_bytes,
                "docx_filename": latest_docx.file_name if latest_docx else None,
                "pdf_filename": latest_pdf.file_name if latest_pdf else None,
                "source": "latest_version",
            }

    generated_docx_bytes, generated_pdf_bytes = await _render_document_files(db, document)
    if generated_docx_bytes or generated_pdf_bytes:
        return {
            "docx_bytes": generated_docx_bytes,
            "pdf_bytes": generated_pdf_bytes,
            "docx_filename": None,
            "pdf_filename": None,
            "source": "generated_from_body",
        }

    return {
        "docx_bytes": current_docx_bytes,
        "pdf_bytes": current_pdf_bytes,
        "docx_filename": current_docx.file_name if current_docx else None,
        "pdf_filename": current_pdf.file_name if current_pdf else None,
        "source": "current_render",
    }


async def _store_current_render_files(
    db: AsyncSession,
    document: OutgoingDocument,
    docx_bytes: bytes,
    pdf_bytes: bytes,
) -> OutgoingDocumentDetailResponse:
    current_files = await _get_current_render_files(db, str(document.id))
    for item in current_files:
        if item.file_path:
            try:
                await delete_path(item.file_path)
            except Exception:
                pass
        await db.delete(item)
    await db.flush()

    current_paths = _build_current_render_paths(document)
    await ensure_path(current_paths["root"])
    await upload_bytes_with_safe_extension(current_paths["docx_path"], docx_bytes)
    await upload_bytes_with_safe_extension(current_paths["pdf_path"], pdf_bytes)
    docx_url = await publish(current_paths["docx_path"])
    pdf_url = await publish(current_paths["pdf_path"])

    db.add(
        OutgoingDocumentFile(
            document_id=str(document.id),
            file_type="render_docx",
            file_path=current_paths["docx_path"],
            file_name=current_paths["docx_name"],
            public_url=docx_url,
        )
    )
    db.add(
        OutgoingDocumentFile(
            document_id=str(document.id),
            file_type="render_pdf",
            file_path=current_paths["pdf_path"],
            file_name=current_paths["pdf_name"],
            public_url=pdf_url,
        )
    )
    await db.commit()
    return await _serialize_document(db, document, include_details=True)


def _format_date_long_ru(value: date) -> str:
    months = [
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
    ]
    return f"{value.day:02d} {months[value.month - 1]} {value.year}"


def _build_outgoing_file_base_clean(document: OutgoingDocument) -> str:
    label = _company_label(document.our_company_key).replace(" ", "_")
    display_number = _display_outgoing_number(document.outgoing_number) or document.outgoing_number
    number_part = _format_outgoing_number_for_file(display_number)
    kind_label = DOCUMENT_KIND_LABELS.get(_normalize_document_kind(getattr(document, "document_kind", None)), "Документ")
    return f"{label}_{kind_label}_{number_part}"


def _ascii_filename_fallback(filename: str, default_stem: str = "document") -> str:
    value = unicodedata.normalize("NFKD", filename or "").encode("ascii", "ignore").decode("ascii")
    value = re.sub(r"[^\w\-.() ]+", "_", value).strip().strip(".")
    if not value:
        value = default_stem
    if "." not in value:
        return value
    stem, ext = value.rsplit(".", 1)
    return f"{stem or default_stem}.{ext}"


def _content_disposition(disposition: str, filename: str) -> str:
    fallback = _ascii_filename_fallback(filename)
    encoded = quote(filename, safe="")
    return f"{disposition}; filename=\"{fallback}\"; filename*=UTF-8''{encoded}"


def _render_pdf_html(
    outgoing_number: str,
    letter_date: date,
    recipient_name: str,
    recipient_short_name: str,
    recipient_to_name: str,
    recipient_appeal: str,
    recipient_eio: str,
    recipient_salutation: str,
    subject: str,
    body: str,
    attachments_list: str,
    company_key: str,
) -> str:
    profile = _company_render_profile(company_key)
    logo_src = _asset_data_uri(profile.get("logo", ""))
    signature_src = _asset_data_uri(profile.get("signature", ""))
    safe_body = _sanitize_html(body)
    safe_attachments = _escape_text(attachments_list).replace("\n", "<br>")
    safe_recipient = _escape_text(recipient_name)
    safe_recipient_short = _escape_text(recipient_short_name)
    safe_recipient_to = _escape_text(recipient_to_name)
    safe_recipient_appeal = _escape_text(recipient_appeal)
    safe_recipient_eio = _escape_text(recipient_eio)
    safe_recipient_salutation = _escape_text(recipient_salutation)
    safe_subject = _escape_text(subject)
    safe_company = _escape_text(_company_label(company_key))
    safe_salutation = " ".join(part for part in [safe_recipient_salutation, safe_recipient_appeal] if part).strip()
    safe_signer_title = _escape_text(profile.get("signer_title", "Руководитель"))
    safe_signer_name = _escape_text(profile.get("signer_name", ""))
    safe_company_head = safe_signer_name
    info_markup = "".join(f"<div>{_escape_text(line)}</div>" for line in profile.get("info_lines", []))
    footer_markup = "".join(f"<div>{_escape_text(line)}</div>" for line in profile.get("footer_lines", []))
    return f"""
<!DOCTYPE html>
<html lang="ru">
  <head>
    <meta charset="utf-8">
    <style>
      body {{ font-family: DejaVu Sans, Arial, sans-serif; font-size: 12px; color: #1d2326; }}
      .header {{ display: flex; justify-content: space-between; margin-bottom: 24px; }}
      .title {{ font-size: 16px; font-weight: bold; }}
      .meta {{ margin-bottom: 16px; }}
      .block {{ margin-bottom: 12px; }}
      .label {{ color: #4a545a; font-weight: bold; }}
      .value {{ margin-top: 4px; }}
      .value table {{ width: 100%; border-collapse: collapse; margin: 12px 0; table-layout: fixed; }}
      .value th, .value td {{ border: 1px solid #cbd5e1; padding: 8px 10px; vertical-align: top; word-wrap: break-word; }}
      .value th {{ background: #f8fafc; font-weight: bold; }}
      .value p {{ margin: 0 0 8px 0; }}
      .value ul, .value ol {{ margin: 8px 0 8px 18px; padding: 0; }}
      .footer {{ margin-top: 32px; }}
    </style>
  </head>
  <body>
    <div class="header">
      <div class="title">{safe_company}</div>
      <div>
        <div class="label">Исходящий №</div>
        <div class="value">{outgoing_number}</div>
        <div class="value">{_format_date_long(letter_date)}</div>
      </div>
    </div>

    <div class="block">
      <div class="label">Кому:</div>
      <div class="value">{safe_recipient_to or safe_recipient}</div>
    </div>

    <div class="block">
      <div class="label">Краткое наименование:</div>
      <div class="value">{safe_recipient_short}</div>
    </div>

    <div class="block">
      <div class="label">Обращение:</div>
      <div class="value">{(safe_recipient_salutation + ' ' + safe_recipient_appeal).strip()}</div>
    </div>

    <div class="block">
      <div class="label">ЕИО:</div>
      <div class="value">{safe_recipient_eio}</div>
    </div>

    <div class="block">
      <div class="label">Тема:</div>
      <div class="value">{safe_subject}</div>
    </div>

    <div class="block">
      <div class="label">Текст письма:</div>
      <div class="value">{safe_body}</div>
    </div>

    <div class="block">
      <div class="label">Приложения:</div>
      <div class="value">{safe_attachments}</div>
    </div>

    <div class="footer">
      <div class="label">Руководитель:</div>
      <div class="value">{safe_company_head}</div>
    </div>
  </body>
</html>
"""


def _html_to_pdf(html: str) -> bytes:
    try:
        from xhtml2pdf import pisa
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="PDF generator is not installed") from exc
    def _block_external_resources(uri, rel):
        # H1 defence-in-depth: refuse to resolve any external/local resource
        # (file://, http(s)://, network paths). Only inline data: URIs pass.
        if uri and str(uri).strip().lower().startswith("data:"):
            return uri
        raise RuntimeError("External resource loading is disabled in PDF rendering")

    output = BytesIO()
    result = pisa.CreatePDF(
        html, dest=output, encoding="UTF-8", link_callback=_block_external_resources
    )
    if result.err:
        raise HTTPException(status_code=500, detail="Failed to generate PDF")
    return output.getvalue()


def _render_company_html(
    outgoing_number: str,
    letter_date: date,
    recipient_name: str,
    recipient_short_name: str,
    recipient_to_name: str,
    recipient_appeal: str,
    recipient_eio: str,
    recipient_salutation: str,
    subject: str,
    body: str,
    attachments_list: str,
    company_key: str,
) -> str:
    profile = _company_render_profile(company_key)
    logo_src = _asset_data_uri(profile.get("logo", ""))
    signature_src = _asset_data_uri(profile.get("signature", ""))
    safe_body = _sanitize_html(body)
    safe_attachments = _escape_text(attachments_list).replace("\n", "<br>")
    safe_recipient = _escape_text(recipient_name)
    safe_recipient_short = _escape_text(recipient_short_name)
    safe_recipient_to = _escape_text(recipient_to_name)
    safe_recipient_eio = _escape_text(recipient_eio)
    safe_subject = _escape_text(subject)
    safe_company = _escape_text(_company_label(company_key))
    safe_salutation = _escape_text(" ".join(part for part in [recipient_salutation, recipient_appeal] if part).strip())
    safe_signer_title = _escape_text(profile.get("signer_title", "Руководитель"))
    safe_signer_name = _escape_text(profile.get("signer_name", ""))
    info_markup = "".join(f"<div>{_escape_text(line)}</div>" for line in profile.get("info_lines", []))
    footer_markup = "".join(f"<div>{_escape_text(line)}</div>" for line in profile.get("footer_lines", []))
    attachments_markup = (
        f"""
        <div class="section">
          <div class="section-label">Приложение:</div>
          <div class="section-value">{safe_attachments}</div>
        </div>
        """
        if safe_attachments
        else ""
    )
    subject_markup = f"<div class=\"doc-meta-subject\">[{safe_subject}]</div>" if safe_subject else ""
    return f"""
<!DOCTYPE html>
<html lang="ru">
  <head>
    <meta charset="utf-8">
    <style>
      @page {{ size: A4; margin: 26mm 18mm 24mm 18mm; }}
      body {{ font-family: "Times New Roman", serif; font-size: 14px; color: #111111; margin: 0; }}
      .page {{ width: 100%; }}
      .brand-table, .dochead-table, .signature-table {{ width: 100%; border-collapse: collapse; }}
      .brand-table td, .dochead-table td, .signature-table td {{ vertical-align: top; }}
      .brand-left {{ width: 44%; }}
      .brand-logo {{ max-width: 250px; max-height: 78px; }}
      .brand-company {{ font-size: 28px; font-weight: bold; color: #cb1517; letter-spacing: 0.6px; }}
      .brand-right {{ font-size: 12px; line-height: 1.35; text-align: right; }}
      .brand-line {{ border-top: 4px solid #cb1517; margin: 10px 0 22px; }}
      .dochead-meta {{ width: 48%; font-size: 14px; line-height: 1.35; }}
      .dochead-recipient {{ width: 42%; font-size: 14px; line-height: 1.35; text-align: right; }}
      .doc-meta-subject {{ margin-top: 6px; }}
      .salutation {{ margin: 28px 0 18px; text-align: center; font-size: 15px; }}
      .body {{ text-align: justify; line-height: 1.4; }}
      .body p {{ margin: 0 0 10px 0; text-indent: 1.25cm; }}
      .body ul, .body ol, .section-value ul, .section-value ol {{ margin: 8px 0 8px 22px; padding: 0; }}
      .body table, .section-value table {{ width: 100%; border-collapse: collapse; margin: 12px 0; table-layout: fixed; }}
      .body th, .body td, .section-value th, .section-value td {{ border: 1px solid #cbd5e1; padding: 7px 9px; vertical-align: top; word-wrap: break-word; }}
      .body th, .section-value th {{ background: #f8fafc; font-weight: bold; text-align: center; }}
      .section {{ margin-top: 18px; }}
      .section-label {{ font-weight: bold; margin-bottom: 6px; }}
      .signature-table {{ margin-top: 38px; }}
      .signature-title {{ padding-bottom: 10px; }}
      .signature-image {{ width: 115px; height: auto; }}
      .signature-name {{ text-align: right; width: 34%; }}
      .footer {{ margin-top: 42px; padding-top: 14px; border-top: 1px solid #d8e0f0; font-size: 11px; line-height: 1.35; text-align: center; }}
    </style>
  </head>
  <body>
    <div class="page">
      <table class="brand-table">
        <tr>
          <td class="brand-left">
            {f'<img class="brand-logo" src="{logo_src}" alt="{safe_company}">' if logo_src else f'<div class="brand-company">{safe_company}</div>'}
          </td>
          <td class="brand-right">{info_markup}</td>
        </tr>
      </table>
      <div class="brand-line"></div>

      <table class="dochead-table">
        <tr>
          <td class="dochead-meta">
            <div>Исх. № {outgoing_number}</div>
            <div>от {_format_date_long_ru(letter_date)} г.</div>
            {subject_markup}
          </td>
          <td class="dochead-recipient">
            {f'<div>{safe_recipient_short}</div>' if safe_recipient_short else ''}
            <div>{safe_recipient_to or safe_recipient}</div>
            {f'<div>{safe_recipient_eio}</div>' if safe_recipient_eio else ''}
          </td>
        </tr>
      </table>

      {f'<div class="salutation">{safe_salutation}!</div>' if safe_salutation else ''}
      <div class="body">{safe_body or '<p>Текст отсутствует</p>'}</div>
      {attachments_markup}

      <table class="signature-table">
        <tr>
          <td><div class="signature-title">{safe_signer_title}</div></td>
          <td class="signature-name"></td>
        </tr>
        <tr>
          <td>{f'<img class="signature-image" src="{signature_src}" alt="{safe_signer_name}">' if signature_src else ''}</td>
          <td class="signature-name">{safe_signer_name}</td>
        </tr>
      </table>

      {f'<div class="footer">{footer_markup}</div>' if footer_markup else ''}
    </div>
  </body>
</html>
"""


_DOC_CONVERT_TIMEOUT = 90  # seconds — LibreOffice headless conversion
_NODE_RENDER_TIMEOUT = 120  # seconds — Node docx renderer


def _convert_docx_to_pdf_bytes(data: bytes) -> Optional[bytes]:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        docx_path = tmp_path / "input.docx"
        docx_path.write_bytes(data)
        try:
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--norestore",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_path),
                    str(docx_path),
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=_DOC_CONVERT_TIMEOUT,
            )
        except subprocess.TimeoutExpired:
            return None
        pdf_path = tmp_path / "input.pdf"
        if result.returncode == 0 and pdf_path.exists():
            return pdf_path.read_bytes()
        return None


def _render_docx_from_html_bytes(html: str) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        html_path = tmp_path / "letter.html"
        html_path.write_text(html, encoding="utf-8")
        odt_path = tmp_path / "letter.odt"
        docx_path = tmp_path / "letter.docx"

        try:
            html_to_odt = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--norestore",
                    "--convert-to",
                    "odt",
                    "--outdir",
                    str(tmp_path),
                    str(html_path),
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=_DOC_CONVERT_TIMEOUT,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError("HTML to ODT conversion timed out")
        if html_to_odt.returncode != 0 or not odt_path.exists():
            raise RuntimeError("Failed to convert HTML to ODT")

        try:
            odt_to_docx = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--norestore",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(tmp_path),
                    str(odt_path),
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=_DOC_CONVERT_TIMEOUT,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError("ODT to DOCX conversion timed out")
        if odt_to_docx.returncode != 0 or not docx_path.exists():
            raise RuntimeError("Failed to convert ODT to DOCX")
        return docx_path.read_bytes()


def _render_docx_from_html_v2(html: str) -> bytes:
    """Pure-Python HTML -> DOCX for template_v2 documents.

    LibreOffice (soffice) is not available on every deployment, so the v2
    pipeline cannot rely on the HTML->ODT->DOCX path. This converter handles
    the subset of HTML emitted by ``_render_v2_html_document``: headings,
    paragraphs, inline bold/italic/underline, line breaks, simple lists and
    tables, plus block-level text alignment.
    """
    from io import BytesIO
    from bs4 import BeautifulSoup, NavigableString, Tag
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    soup = BeautifulSoup(html or "", "lxml")
    body = soup.body or soup
    doc = DocxDocument()
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(13)
        npf = normal.paragraph_format
        npf.space_before = Pt(0)
        npf.space_after = Pt(0)
        npf.line_spacing = 1.0
    except Exception:
        pass

    # A4 + reference letter margins (ГОСТ-like: L 2.5, R 1.5, T/B 1.5 cm)
    sec = doc.sections[0]
    try:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.5)
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
    except Exception:
        pass
    CONTENT_CM = 21.0 - 2.5 - 1.5  # usable text width

    BOLD_TAGS = {"b", "strong"}
    ITALIC_TAGS = {"i", "em"}
    UNDERLINE_TAGS = {"u", "ins"}
    BLOCK_TAGS = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "div",
                  "table", "ul", "ol", "section", "article"}
    RIGHT_CLASSES = ("text-right", "doc-recipient", "doc-sign-name",
                     "doc-letter__recipient")
    CENTER_CLASSES = ("text-center", "doc-salutation", "doc-center", "doc-title")
    JUSTIFY_CLASSES = ("text-justify", "doc-body", "doc-justify")
    INDENT_CLASSES = ("doc-body", "doc-indent")

    def _classes(node):
        return set(node.get("class") or []) if isinstance(node, Tag) else set()

    def _align_of(node):
        cls = _classes(node)
        st = (node.get("style") or "").lower() if isinstance(node, Tag) else ""
        if cls & set(RIGHT_CLASSES) or "text-align:right" in st or "text-align: right" in st:
            return WD_ALIGN_PARAGRAPH.RIGHT
        if cls & set(CENTER_CLASSES) or "text-align:center" in st or "text-align: center" in st:
            return WD_ALIGN_PARAGRAPH.CENTER
        if cls & set(JUSTIFY_CLASSES) or "justify" in st:
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return None

    def _indent_of(node):
        """First-line indent (cm) for letter body paragraphs."""
        cls = _classes(node)
        st = (node.get("style") or "").lower() if isinstance(node, Tag) else ""
        if cls & set(INDENT_CLASSES):
            return Cm(1.25)
        m = re.search(r"text-indent:\s*([\d.]+)cm", st)
        if m:
            try:
                return Cm(float(m.group(1)))
            except ValueError:
                return None
        return None

    def _table_is_borderless(node):
        cls = _classes(node)
        st = (node.get("style") or "").lower() if isinstance(node, Tag) else ""
        if {"nb", "borderless", "no-border"} & cls:
            return True
        if "border:0" in st or "border: 0" in st or "border:none" in st or "border: none" in st:
            return True
        return False

    def _col_pct(node):
        st = (node.get("style") or "").lower() if isinstance(node, Tag) else ""
        m = re.search(r"width:\s*([\d.]+)%", st)
        if m:
            try:
                return float(m.group(1))
            except ValueError:
                return None
        return None

    def _inline(node, paragraph, fmt):
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                run = paragraph.add_run(text)
                run.bold = True if fmt.get("b") else None
                run.italic = True if fmt.get("i") else None
                run.underline = True if fmt.get("u") else None
            return
        if not isinstance(node, Tag):
            return
        name = (node.name or "").lower()
        if name == "br":
            paragraph.add_run().add_break()
            return
        nf = dict(fmt)
        if name in BOLD_TAGS:
            nf["b"] = True
        if name in ITALIC_TAGS:
            nf["i"] = True
        if name in UNDERLINE_TAGS:
            nf["u"] = True
        st = (node.get("style") or "").lower()
        if "font-weight:bold" in st or "font-weight: bold" in st or "font-weight:700" in st:
            nf["b"] = True
        if "font-style:italic" in st or "font-style: italic" in st:
            nf["i"] = True
        if "underline" in st:
            nf["u"] = True
        for child in node.children:
            _inline(child, paragraph, nf)

    def _add_paragraph(node, style=None, force_align=None, force_indent=None):
        try:
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        except Exception:
            p = doc.add_paragraph()
        al = _align_of(node) or force_align
        if al is not None:
            p.alignment = al
        ind = _indent_of(node) or force_indent
        if ind is not None:
            p.paragraph_format.first_line_indent = ind
        for child in node.children:
            _inline(child, p, {})
        return p

    def _add_table(node):
        rows = node.find_all("tr")
        if not rows:
            return
        ncols = max((len(r.find_all(["td", "th"])) for r in rows), default=0)
        if ncols <= 0:
            return
        borderless = _table_is_borderless(node)
        table = doc.add_table(rows=0, cols=ncols)
        table.autofit = False
        if not borderless:
            try:
                table.style = "Table Grid"
            except Exception:
                pass
        # Column widths from first row's <td style="width:NN%">
        first_cells = rows[0].find_all(["td", "th"])
        pcts = [_col_pct(c) for c in first_cells]
        widths = None
        if any(pcts):
            filled = [pp if pp else (100.0 / ncols) for pp in pcts]
            total = sum(filled) or 100.0
            widths = [Cm(CONTENT_CM * pp / total) for pp in filled]
        for r in rows:
            cells = r.find_all(["td", "th"])
            row_cells = table.add_row().cells
            for idx in range(ncols):
                cell = row_cells[idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                if widths and idx < len(widths):
                    cell.width = widths[idx]
                para = cell.paragraphs[0]
                para.text = ""
                if idx < len(cells):
                    src = cells[idx]
                    cal = _align_of(src)
                    if cal is not None:
                        para.alignment = cal
                    base = {"b": True} if (src.name or "").lower() == "th" else {}
                    for ch in src.children:
                        _inline(ch, para, base)
        if borderless:
            tbl_pr = table._tbl.tblPr
            borders = OxmlElement("w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                e = OxmlElement(f"w:{edge}")
                e.set(qn("w:val"), "none")
                e.set(qn("w:sz"), "0")
                e.set(qn("w:space"), "0")
                borders.append(e)
            tbl_pr.append(borders)

    def _walk(parent, inherited_align=None, inherited_indent=None):
        for node in parent.children:
            if isinstance(node, NavigableString):
                txt = str(node).strip()
                if txt:
                    p = doc.add_paragraph()
                    if inherited_align is not None:
                        p.alignment = inherited_align
                    if inherited_indent is not None:
                        p.paragraph_format.first_line_indent = inherited_indent
                    p.add_run(txt)
                continue
            if not isinstance(node, Tag):
                continue
            name = (node.name or "").lower()
            if name in ("script", "style"):
                continue
            if name == "hr" or "doc-rule" in _classes(node):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                pPr = p._p.get_or_add_pPr()
                pbdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "18")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "000000")
                pbdr.append(bottom)
                pPr.append(pbdr)
                continue
            if name == "table":
                _add_table(node)
            elif name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = min(int(name[1]), 4)
                _add_paragraph(node, style=f"Heading {level}",
                               force_align=inherited_align)
            elif name == "p":
                _add_paragraph(node, force_align=inherited_align,
                               force_indent=inherited_indent)
            elif name in ("ul", "ol"):
                bullet = "List Bullet" if name == "ul" else "List Number"
                for li in node.find_all("li", recursive=False):
                    try:
                        p = doc.add_paragraph(style=bullet)
                    except Exception:
                        p = doc.add_paragraph()
                    for ch in li.children:
                        _inline(ch, p, {})
            elif name in ("div", "section", "article", "body"):
                child_align = _align_of(node) or inherited_align
                child_indent = _indent_of(node) or inherited_indent
                has_block = any(
                    isinstance(c, Tag) and (c.name or "").lower() in BLOCK_TAGS
                    for c in node.children
                )
                if has_block:
                    _walk(node, child_align, child_indent)
                else:
                    _add_paragraph(node, force_align=child_align,
                                   force_indent=child_indent)
            else:
                _walk(node, inherited_align, inherited_indent)

    _walk(body)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def _build_document_render_html(db: AsyncSession, document: OutgoingDocument) -> str:
    render_payload = await _build_document_render_payload(db, document)
    recipient_payload = render_payload.get("recipient") or {}
    document_payload = render_payload.get("document") or {}
    return _render_company_html(
        outgoing_number=render_payload.get("outgoing_number") or _display_outgoing_number(document.outgoing_number) or document.outgoing_number,
        letter_date=document.letter_date,
        recipient_name=recipient_payload.get("name") or "",
        recipient_short_name=recipient_payload.get("short_name") or "",
        recipient_to_name=recipient_payload.get("to_name") or "",
        recipient_appeal=recipient_payload.get("appeal") or "",
        recipient_eio=recipient_payload.get("eio") or "",
        recipient_salutation=recipient_payload.get("salutation") or "",
        subject=document_payload.get("subject") or "",
        body=render_payload.get("body") or "",
        attachments_list=render_payload.get("attachments_list") or "",
        company_key=document.our_company_key or "normbud",
    )


async def _build_document_render_payload(db: AsyncSession, document: OutgoingDocument) -> dict:
    recipient = await Company.get_by_id(db, str(document.recipient_company_id))
    if not recipient:
        raise HTTPException(status_code=404, detail="Recipient company not found")
    template_version = await _get_active_document_template_version(db, document)
    deal = await Deal.get_by_id(db, str(document.deal_id)) if document.deal_id else None
    contract = await Contract.get_by_id(db, str(document.contract_id)) if document.contract_id else None
    company_key = _normalize_company_key(document.our_company_key)
    our_company_model = await _load_our_company_for_document(db, company_key, deal)

    recipient_bank = _json_loads_or_default(getattr(document, "bank_account_snapshot", None), None)
    if not recipient_bank:
        recipient_accounts = _recipient_bank_accounts(recipient)
        recipient_bank = recipient_accounts[0] if recipient_accounts else None
    our_accounts = _recipient_bank_accounts(our_company_model) if our_company_model else []
    our_bank = our_accounts[0] if our_accounts else None

    linked_stage_ids = _json_loads_or_default(getattr(document, "linked_stage_ids", None), []) or []
    linked_stages = []
    for stage_id in linked_stage_ids:
        stage = await Stage.get_by_id(db, str(stage_id))
        if stage:
            linked_stages.append(stage)

    payment_items = _json_loads_or_default(getattr(document, "linked_payment_items", None), []) or []
    if not isinstance(payment_items, list):
        payment_items = []

    payment_total = sum(_payment_item_amount(item) for item in payment_items)
    stage_total = sum(_safe_float(getattr(stage, "planned_cost", 0)) for stage in linked_stages)
    contract_total = _safe_float(getattr(contract, "amount", 0))
    source_amount = payment_total or stage_total or contract_total
    vat_rate = _safe_float(getattr(deal, "vat_rate", 20)) if deal else 20.0
    vat_included = bool(getattr(deal, "vat_included", True)) if deal else True
    if source_amount and vat_rate > 0:
        if vat_included:
            total_amount = source_amount
            vat_amount = source_amount * vat_rate / (100 + vat_rate)
            amount_without_vat = source_amount - vat_amount
        else:
            vat_amount = source_amount * vat_rate / 100
            total_amount = source_amount + vat_amount
            amount_without_vat = source_amount
    else:
        total_amount = source_amount
        vat_amount = 0.0
        amount_without_vat = source_amount

    display_number = _display_outgoing_number(document.outgoing_number) or document.outgoing_number or ""
    document_date = _format_ru_date(document.letter_date)
    contract_number = _first_value(getattr(contract, "contract_number", None))
    contract_date = _format_ru_date(getattr(contract, "contract_date", None))
    document_basis = ""
    if contract_number or contract_date:
        number_part = _contract_number_with_prefix(contract_number) if contract_number else ""
        date_part = f" от {contract_date}" if contract_date else ""
        document_basis = f"Договор {number_part}{date_part}".strip()

    recipient_payload = _company_payload(recipient, bank_account=recipient_bank)
    recipient_payload.update({
        "short_name": _first_value(document.recipient_short_name, recipient_payload.get("short_name")),
        "to_name": document.recipient_to_name or "",
        "genitive_name": document.recipient_genitive_name or "",
        "appeal": document.recipient_appeal or "",
        "eio": document.recipient_eio or "",
        "salutation": document.recipient_salutation or "",
    })

    company_profile = _company_render_profile(company_key)
    director_name = _director_short_name(company_profile.get("signer_name"))
    our_company_payload = _company_payload(
        our_company_model,
        fallback=OUR_COMPANY_REQUISITES.get(company_key, {}),
        bank_account=our_bank,
        director_name=director_name,
    )
    our_company_payload["signer_title"] = company_profile.get("signer_title", "")
    our_company_payload["signer_name"] = company_profile.get("signer_name", "")

    payment_due_date = _add_workdays(document.letter_date, 5)
    vat_rate_value = round(float(vat_rate or 0), 2)
    vat_rate_text = _format_rate(vat_rate_value)

    document_payload = {
        "id": str(document.id),
        "kind": _normalize_document_kind(document.document_kind),
        "kind_label": DOCUMENT_KIND_LABELS.get(_normalize_document_kind(document.document_kind), ""),
        "number": display_number,
        "date": document_date,
        "date_iso": document.letter_date.isoformat() if document.letter_date else "",
        "subject": document.subject or "",
        "basis": document_basis,
        "total_amount": _format_money(total_amount),
        "total_amount_raw": round(float(total_amount or 0), 2),
        "total_amount_words": _amount_to_words_ru(total_amount),
        "amount_without_vat": _format_money(amount_without_vat),
        "amount_without_vat_raw": round(float(amount_without_vat or 0), 2),
        "vat_amount": _format_money(vat_amount),
        "vat_amount_raw": round(float(vat_amount or 0), 2),
        "vat_rate": vat_rate_value,
        "vat_rate_text": vat_rate_text,
        "payment_due_date": _format_ru_date(payment_due_date),
        "payment_due_date_iso": payment_due_date.isoformat() if payment_due_date else "",
    }

    contract_payload = {
        "id": str(contract.id) if contract else "",
        "number": contract_number,
        "date": contract_date,
        "amount": _format_money(contract_total),
        "amount_raw": round(float(contract_total or 0), 2),
    }

    stage_payloads = [
        {
            "id": str(stage.id),
            "name": stage.name or "",
            "amount": _format_money(getattr(stage, "planned_cost", 0)),
            "amount_raw": round(_safe_float(getattr(stage, "planned_cost", 0)), 2),
        }
        for stage in linked_stages
    ]
    payment_payloads = [
        {
            "entry_id": str(item.get("entry_id") or item.get("payment_id") or ""),
            "amount": _format_money(_payment_item_amount(item)),
            "amount_raw": _payment_item_amount(item),
            "note": str(item.get("note") or ""),
        }
        for item in payment_items
        if isinstance(item, dict)
    ]

    invoice_payload = {
        "number": document_payload["number"],
        "date": document_payload["date"],
        "basis": document_payload["basis"],
        "total_amount": document_payload["total_amount"],
        "total_amount_raw": document_payload["total_amount_raw"],
        "total_amount_words": document_payload["total_amount_words"],
        "vat_amount": document_payload["vat_amount"],
        "vat_amount_raw": document_payload["vat_amount_raw"],
        "vat_rate": document_payload["vat_rate"],
        "vat_rate_text": document_payload["vat_rate_text"],
        "payment_due_date": document_payload["payment_due_date"],
        "payment_due_date_iso": document_payload["payment_due_date_iso"],
    }

    render_payload = {
        "id": str(document.id),
        "outgoing_number": display_number,
        "our_company_key": company_key,
        "template_file_path": template_version.file_path if template_version else "",
        "letter_date": document.letter_date.isoformat() if document.letter_date else "",
        "subject": document.subject or "",
        "body": document.body or "",
        "attachments_list": document.attachments_list or "",
        "recipient_company_name": recipient.name or "",
        "recipient_short_name": recipient_payload.get("short_name") or "",
        "recipient_to_name": document.recipient_to_name or "",
        "recipient_appeal": document.recipient_appeal or "",
        "recipient_eio": document.recipient_eio or "",
        "recipient_salutation": document.recipient_salutation or "",
        "recipient_genitive_name": document.recipient_genitive_name or "",
        "signer_title": company_profile.get("signer_title", ""),
        "signer_name": company_profile.get("signer_name", ""),
        "document": document_payload,
        "invoice": invoice_payload,
        "deal": {
            "id": str(deal.id) if deal else "",
            "title": getattr(deal, "title", "") or "",
            "obj_name": getattr(deal, "obj_name", "") or "",
            "address": getattr(deal, "address", "") or "",
            "vat_rate": round(float(vat_rate or 0), 2),
            "vat_included": vat_included,
        },
        "contract": contract_payload,
        "recipient": recipient_payload,
        "our_company": our_company_payload,
        "stages": stage_payloads,
        "linked_payment_items": payment_payloads,
    }
    editor_mode = normalize_editor_mode(getattr(document, "editor_mode", None))
    editor_draft = normalize_editor_draft(
        _json_loads_or_default(getattr(document, "editor_draft_json", None), None),
        document_payload["kind"],
    )
    if editor_mode == "structured":
        render_payload = _apply_structured_editor_overrides(document, render_payload, editor_draft)
    else:
        render_payload["editor"] = {
            "mode": editor_mode,
            "schema_version": int(getattr(document, "editor_schema_version", None) or 1),
            "draft": editor_draft,
            "block_types": [block.get("type") for block in get_editor_blocks(editor_draft, document_payload["kind"])],
            "resolved": {
                "basis": document_payload.get("basis") or "",
                "payment_due_date": document_payload.get("payment_due_date") or "",
                "body_html": render_payload.get("body") or "",
            },
        }
    return render_payload


async def _latest_template_version(db: AsyncSession, template: DocumentTemplate) -> Optional[DocumentTemplateVersion]:
    if template.current_version_id:
        result = await db.execute(
            select(DocumentTemplateVersion).where(DocumentTemplateVersion.id == str(template.current_version_id))
        )
        version = result.scalar_one_or_none()
        if version:
            return version
    result = await db.execute(
        select(DocumentTemplateVersion)
        .where(DocumentTemplateVersion.template_id == str(template.id))
        .order_by(DocumentTemplateVersion.version_number.desc())
        .limit(1)
    )
    return result.scalar_one_or_none()


async def _get_active_document_template_version(
    db: AsyncSession,
    document: OutgoingDocument,
) -> Optional[DocumentTemplateVersion]:
    document_kind = _normalize_document_kind(getattr(document, "document_kind", None))
    company_key = _normalize_company_key(getattr(document, "our_company_key", None))
    scopes = [
        DocumentTemplate.our_company_key == company_key,
        DocumentTemplate.our_company_key.is_(None),
    ]
    for scope in scopes:
        result = await db.execute(
            select(DocumentTemplate)
            .where(
                and_(
                    DocumentTemplate.module == "outgoing_registry",
                    DocumentTemplate.document_kind == document_kind,
                    DocumentTemplate.is_active == True,  # noqa: E712
                    DocumentTemplate.status != "archived",
                    scope,
                )
            )
            .order_by(DocumentTemplate.updated_at.desc(), DocumentTemplate.created_at.desc())
            .limit(5)
        )
        for template in result.scalars().all():
            version = await _latest_template_version(db, template)
            if version and version.file_path:
                return version
    return None


def _render_docx_via_node(payload: dict) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        input_path = tmp_path / "outgoing.json"
        output_path = tmp_path / "outgoing.docx"
        input_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
        try:
            result = subprocess.run(
                ["node", str(RENDER_SCRIPT_PATH), str(input_path), str(output_path)],
                cwd=str(FRONTEND_DIR),
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=_NODE_RENDER_TIMEOUT,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError("Node DOCX renderer timed out")
        if result.returncode != 0 or not output_path.exists():
            raise RuntimeError(
                "Node DOCX renderer failed:\n"
                f"stdout:\n{result.stdout}\n"
                f"stderr:\n{result.stderr}"
            )
        return output_path.read_bytes()


def _render_v2_html_document(
    layout_html: str,
    field_values: dict,
    body_html: str,
    stages: list,
    payments: list,
) -> str:
    """Compose the final HTML for a template_v2 document.

    Substitutes [data-placeholder] with resolved values, injects the editable
    body into [data-editable], and expands [data-table="stages|payments"] into
    real tables. Output is a self-contained HTML page handed to the existing
    HTML→DOCX / HTML→PDF helpers.
    """
    from bs4 import BeautifulSoup

    # H1: the editable body is user-authored — sanitize before it reaches the
    # HTML→PDF (xhtml2pdf) / HTML→DOCX paths to block file:// LFR and SSRF.
    body_html = _sanitize_html(body_html or "")

    soup = BeautifulSoup(layout_html or "", "lxml")
    root = soup.body or soup

    for el in root.select("[data-placeholder]"):
        key = el.get("data-placeholder") or ""
        val = field_values.get(key)
        el.clear()
        el.append("" if val is None else str(val))
        del el["data-placeholder"]

    for el in root.select("[data-editable]"):
        frag = BeautifulSoup(body_html or "<p></p>", "lxml")
        el.clear()
        for child in list((frag.body or frag).children):
            el.append(child)
        del el["data-editable"]

    for el in root.select("[data-table]"):
        kind = el.get("data-table")
        rows = stages if kind == "stages" else payments
        head = ("№", "Наименование", "Сумма") if kind == "stages" else ("№", "Назначение", "Сумма")
        table = soup.new_tag("table")
        table["class"] = "doc-table"
        thead = soup.new_tag("thead")
        trh = soup.new_tag("tr")
        for c in head:
            th = soup.new_tag("th"); th.string = c; trh.append(th)
        thead.append(trh); table.append(thead)
        tbody = soup.new_tag("tbody")
        for i, r in enumerate(rows or []):
            tr = soup.new_tag("tr")
            c1 = soup.new_tag("td"); c1.string = str(i + 1)
            label = (r.get("name") if kind == "stages" else (r.get("note") or "Платёж")) or "—"
            c2 = soup.new_tag("td"); c2.string = str(label)
            c3 = soup.new_tag("td"); c3["class"] = "amount"; c3.string = str(r.get("amount") or "—")
            tr.extend([c1, c2, c3]); tbody.append(tr)
        if not (rows or []):
            tr = soup.new_tag("tr")
            td = soup.new_tag("td"); td["colspan"] = "3"
            td.string = "Этапы не выбраны" if kind == "stages" else "Платежи не выбраны"
            tr.append(td); tbody.append(tr)
        table.append(tbody)
        el.replace_with(table)

    inner = "".join(str(c) for c in (root.children if root else []))
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'><style>"
        "@page{size:A4;margin:1.5cm 1.5cm 1.5cm 2.5cm;}"
        "body{font-family:'Times New Roman',serif;font-size:13pt;color:#000;margin:0;padding:0;line-height:1.15;}"
        "p{margin:0 0 6pt;}"
        "h1{font-size:14pt;font-weight:700;text-align:center;text-transform:uppercase;margin:0 0 16pt;}"
        ".doc-title{text-align:center;font-weight:700;text-transform:uppercase;margin:0 0 16pt;}"
        # Letterhead (our-company block at the very top)
        ".doc-letterhead{margin:0 0 6pt;}"
        ".doc-letterhead .lh-name{font-size:15pt;font-weight:700;}"
        ".doc-letterhead .lh-req{font-size:10pt;color:#000;}"
        ".doc-rule{border:0;border-top:2px solid #000;margin:4pt 0 14pt;}"
        # Dochead (Исх № / recipient) — borderless two-column table
        "table.nb,table.nb tr,table.nb td,table.nb th{border:none!important;}"
        "table.nb{width:100%;border-collapse:collapse;margin:0 0 16pt;}"
        "table.nb td{padding:0 4pt;vertical-align:top;font-size:13pt;}"
        ".doc-recipient{text-align:right;}"
        ".doc-letter__recipient{text-align:right;margin-bottom:14pt;}"
        ".doc-salutation{text-align:center;margin:14pt 0 12pt;}"
        ".doc-body{text-align:justify;}"
        ".doc-body p{text-indent:1.25cm;margin:0 0 6pt;text-align:justify;}"
        ".doc-attach{margin:12pt 0 0;}"
        ".doc-sign,.doc-letter__signature{margin-top:30pt;}"
        ".doc-sign td{padding:0 4pt;vertical-align:bottom;}"
        ".doc-sign-name{text-align:right;}"
        ".doc-req p,.doc-totals p{margin:0 0 4pt;}"
        "table.doc-table{width:100%;border-collapse:collapse;margin:12pt 0;font-size:11pt;}"
        "table.doc-table th,table.doc-table td{border:1px solid #000;padding:4pt 6pt;text-align:left;vertical-align:top;}"
        "table.doc-table th{font-weight:700;}"
        "table.doc-table td.amount{text-align:right;white-space:nowrap;}"
        "</style></head><body>" + inner + "</body></html>"
    )


async def _render_v2_files(
    db: AsyncSession, document: OutgoingDocument, draft: dict
) -> tuple[Optional[bytes], bytes]:
    render_payload = await _build_document_render_payload(db, document)
    document_kind = _normalize_document_kind(getattr(document, "document_kind", "letter"))
    resolved = _build_editor_resolved_fields(render_payload, document_kind)
    # Saved draft values reflect what the user actually saw/edited — they win.
    field_values = {**(resolved or {}), **(draft.get("field_values") or {})}
    body_html = ((draft.get("editable_regions") or {}).get("body")) or "<p></p>"
    layout_html = ""
    template_id = draft.get("template_id")
    if template_id:
        from app.models import DocumentTemplate
        from sqlalchemy import select as sa_select
        tpl = (await db.execute(
            sa_select(DocumentTemplate).where(DocumentTemplate.id == str(template_id))
        )).scalar_one_or_none()
        layout_html = (tpl.layout_html if tpl else "") or ""
    if not layout_html:
        # No template layout — nothing v2-specific to render; signal fallback.
        return None, None

    html = _render_v2_html_document(
        layout_html, field_values, body_html,
        render_payload.get("stages") or [],
        render_payload.get("linked_payment_items") or [],
    )
    docx_bytes = None
    try:
        docx_bytes = await run_in_threadpool(_render_docx_from_html_v2, html)
    except Exception:
        docx_bytes = None
    if not docx_bytes:
        try:
            docx_bytes = await run_in_threadpool(_render_docx_from_html_bytes, html)
        except Exception:
            docx_bytes = None
    pdf_bytes = None
    if docx_bytes:
        try:
            pdf_bytes = await run_in_threadpool(_convert_docx_to_pdf_bytes, docx_bytes)
        except Exception:
            pdf_bytes = None
    if pdf_bytes is None:
        pdf_bytes = _html_to_pdf(html)
    return docx_bytes, pdf_bytes


async def _render_document_files(db: AsyncSession, document: OutgoingDocument) -> tuple[Optional[bytes], bytes]:
    html = None
    docx_bytes = None

    # template_v2 path: layout_html-driven render. If the document carries a
    # v2 draft (schema_version 2), render from the inline layout instead of
    # the legacy Node/block pipeline.
    try:
        _draft = _json_loads_or_default(getattr(document, "editor_draft_json", None), None)
        if isinstance(_draft, dict) and int(_draft.get("schema_version") or 0) >= 2:
            v2_docx, v2_pdf = await _render_v2_files(db, document, _draft)
            if v2_pdf is not None:
                return v2_docx, v2_pdf
    except Exception:
        pass  # fall through to legacy pipeline

    try:
        payload = await _build_document_render_payload(db, document)
        docx_bytes = await run_in_threadpool(_render_docx_via_node, payload)
    except Exception:
        docx_bytes = None

    if docx_bytes is None:
        html = await _build_document_render_html(db, document)
        try:
            docx_bytes = await run_in_threadpool(_render_docx_from_html_bytes, html)
        except Exception:
            docx_bytes = None

    pdf_bytes = None
    if docx_bytes:
        try:
            pdf_bytes = await run_in_threadpool(_convert_docx_to_pdf_bytes, docx_bytes)
        except Exception:
            pdf_bytes = None
    if pdf_bytes is None:
        if html is None:
            html = await _build_document_render_html(db, document)
        pdf_bytes = _html_to_pdf(html)
    return docx_bytes, pdf_bytes


def _recipient_bank_accounts(company: Company) -> List[dict]:
    raw = getattr(company, "bank_accounts", None) or []
    if isinstance(raw, str):
        raw = _json_loads_or_default(raw, [])
    if not isinstance(raw, list):
        return []
    return [item for item in raw if isinstance(item, dict)]


def _resolve_bank_account_snapshot(company: Company, bank_account_index: Optional[int]) -> tuple[Optional[int], Optional[dict]]:
    accounts = _recipient_bank_accounts(company)
    if not accounts:
        return None, None
    index = 0 if bank_account_index is None else int(bank_account_index)
    if index < 0 or index >= len(accounts):
        raise HTTPException(status_code=400, detail="Invalid bank_account_index")
    return index, accounts[index]


async def _validate_contract_for_document(
    db: AsyncSession,
    contract_id: Optional[str],
    deal_id: Optional[str],
) -> Optional[Contract]:
    if not contract_id:
        return None
    contract = await Contract.get_by_id(db, str(contract_id))
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    if deal_id and contract.deal_id and str(contract.deal_id) != str(deal_id):
        raise HTTPException(status_code=400, detail="Contract does not belong to selected deal")
    return contract


async def _validate_stage_links(
    db: AsyncSession,
    stage_ids,
    deal_id: Optional[str],
) -> List[str]:
    normalized = [str(item) for item in (_json_loads_or_default(stage_ids, []) or []) if item]
    if not normalized:
        return []
    if not deal_id:
        raise HTTPException(status_code=400, detail="Stage links require selected deal")
    allowed = {str(stage.id): stage for stage in await Stage.get_by_deal_id(db, deal_id)}
    allowed.update({str(stage.id).replace("-", ""): stage for stage in allowed.values()})
    result = []
    for stage_id in normalized:
        stage = allowed.get(stage_id) or allowed.get(stage_id.replace("-", ""))
        if not stage:
            raise HTTPException(status_code=400, detail="Stage does not belong to selected deal")
        result.append(str(stage.id))
    return result


async def _existing_payment_usage(
    db: AsyncSession,
    *,
    exclude_document_id: Optional[str] = None,
) -> dict:
    result = await db.execute(
        select(OutgoingDocument.id, OutgoingDocument.linked_payment_items)
        .where(OutgoingDocument.linked_payment_items.isnot(None))
    )
    usage = {}
    for doc_id, payload in result.all():
        if exclude_document_id and str(doc_id) == str(exclude_document_id):
            continue
        items = _json_loads_or_default(payload, [])
        if not isinstance(items, list):
            continue
        for item in items:
            if not isinstance(item, dict):
                continue
            entry_id = str(item.get("entry_id") or item.get("payment_id") or "").strip()
            if not entry_id:
                continue
            try:
                amount = float(item.get("amount") or 0)
            except (TypeError, ValueError):
                amount = 0.0
            if amount <= 0:
                continue
            usage[entry_id] = usage.get(entry_id, 0.0) + amount
    return usage


async def _validate_payment_links(
    db: AsyncSession,
    payment_items,
    deal_id: Optional[str],
    contract_id: Optional[str],
    *,
    exclude_document_id: Optional[str] = None,
) -> List[dict]:
    raw_items = _json_loads_or_default(payment_items, []) or []
    if not isinstance(raw_items, list):
        raise HTTPException(status_code=400, detail="linked_payment_items must be a list")
    if not raw_items:
        return []
    usage = await _existing_payment_usage(db, exclude_document_id=exclude_document_id)
    normalized = []
    for raw in raw_items:
        if not isinstance(raw, dict):
            raise HTTPException(status_code=400, detail="Payment item must be an object")
        entry_id = str(raw.get("entry_id") or raw.get("payment_id") or "").strip()
        if not entry_id:
            raise HTTPException(status_code=400, detail="Payment item entry_id is required")
        try:
            amount = float(raw.get("amount") or 0)
        except (TypeError, ValueError):
            raise HTTPException(status_code=400, detail="Payment item amount is invalid")
        if amount <= 0:
            raise HTTPException(status_code=400, detail="Payment item amount must be positive")
        entry_result = await db.execute(select(IncomeExpenseEntry).where(_id_conditions(IncomeExpenseEntry.id, entry_id)))
        entry = entry_result.scalar_one_or_none()
        if not entry:
            raise HTTPException(status_code=404, detail="Payment entry not found")
        if deal_id and entry.deal_id and str(entry.deal_id) != str(deal_id):
            raise HTTPException(status_code=400, detail="Payment entry does not belong to selected deal")
        if contract_id and entry.contract_id and str(entry.contract_id) != str(contract_id):
            raise HTTPException(status_code=400, detail="Payment entry does not belong to selected contract")
        used = float(usage.get(entry_id, 0.0))
        if used + amount > float(entry.amount or 0) + 0.005:
            raise HTTPException(status_code=400, detail="Payment linked amount exceeds available payment amount")
        usage[entry_id] = used + amount
        normalized.append({
            "entry_id": entry_id,
            "amount": round(amount, 2),
            "note": str(raw.get("note") or "").strip(),
        })
    return normalized


def _payment_item_amount(item) -> float:
    if not isinstance(item, dict):
        return 0.0
    try:
        return round(float(item.get("amount") or 0), 2)
    except (TypeError, ValueError):
        return 0.0


def _closing_document_label(document: OutgoingDocument) -> str:
    document_kind = _normalize_document_kind(getattr(document, "document_kind", None))
    kind_label = DOCUMENT_KIND_LABELS.get(document_kind) or "Документ"
    number = _display_outgoing_number(getattr(document, "outgoing_number", None)) or getattr(document, "outgoing_number", None)
    return f"{kind_label} № {number}" if number else kind_label


async def _contract_number_for_document(db: AsyncSession, contract_id) -> Optional[str]:
    if not contract_id:
        return None
    contract = await Contract.get_by_id(db, str(contract_id))
    return contract.contract_number if contract else None


async def _closing_document_payload(db: AsyncSession, document: OutgoingDocument, amount: Optional[float] = None) -> dict:
    document_kind = _normalize_document_kind(getattr(document, "document_kind", None))
    display_number = _display_outgoing_number(getattr(document, "outgoing_number", None)) or getattr(document, "outgoing_number", None)
    return {
        "id": str(document.id),
        "document_kind": document_kind,
        "document_kind_label": DOCUMENT_KIND_LABELS.get(document_kind),
        "number": display_number,
        "label": _closing_document_label(document),
        "date": document.letter_date.isoformat() if document.letter_date else None,
        "subject": document.subject,
        "status": document.status,
        "amount": round(float(amount), 2) if amount is not None else None,
        "contract_id": str(document.contract_id) if getattr(document, "contract_id", None) else None,
        "contract_number": await _contract_number_for_document(db, getattr(document, "contract_id", None)),
    }


async def _sync_document_registry(db: AsyncSession, document: OutgoingDocument) -> None:
    display_number = _display_outgoing_number(document.outgoing_number)
    document_kind = _normalize_document_kind(getattr(document, "document_kind", None))
    result = await db.execute(
        select(Document).where(
            and_(
                Document.source_type == "outgoing_registry",
                Document.source_id == str(document.id),
            )
        )
    )
    existing = result.scalar_one_or_none()
    payload = {
        "doc_type": DOCUMENT_REGISTRY_TYPES.get(document_kind, "outgoing_letter"),
        "title": document.subject or DOCUMENT_KIND_DEFAULT_SUBJECTS.get(document_kind, "Исходящий документ"),
        "number": display_number or document.outgoing_number,
        "document_date": document.letter_date,
        "status": document.status or "draft",
        "project_id": str(document.deal_id) if document.deal_id else None,
        "counterparty_id": str(document.recipient_company_id) if document.recipient_company_id else None,
        "our_company_id": _normalize_company_key(document.our_company_key),
        "source_type": "outgoing_registry",
        "source_id": str(document.id),
    }
    if existing:
        for key, value in payload.items():
            setattr(existing, key, value)
        await db.commit()
        await db.refresh(existing)
        return
    doc = Document(**payload)
    db.add(doc)
    await db.commit()


async def _serialize_document(
    db: AsyncSession,
    document: OutgoingDocument,
    include_details: bool = False,
) -> OutgoingDocumentDetailResponse:
    recipient = await Company.get_by_id(db, str(document.recipient_company_id))
    deal = await Deal.get_by_id(db, str(document.deal_id)) if document.deal_id else None
    contract = await Contract.get_by_id(db, str(document.contract_id)) if getattr(document, "contract_id", None) else None
    display_number = _display_outgoing_number(document.outgoing_number)
    company_key = document.our_company_key or "normbud"
    document_kind = _normalize_document_kind(getattr(document, "document_kind", None))
    bank_accounts = _recipient_bank_accounts(recipient) if recipient else []
    selected_bank_account = _json_loads_or_default(getattr(document, "bank_account_snapshot", None), None)
    linked_stage_ids = _json_loads_or_default(getattr(document, "linked_stage_ids", None), [])
    linked_payment_items = _json_loads_or_default(getattr(document, "linked_payment_items", None), [])
    editor_draft = _json_loads_or_default(getattr(document, "editor_draft_json", None), None)
    editor_validation = _json_loads_or_default(getattr(document, "editor_validation_json", None), None)
    editor_render_context = _json_loads_or_default(getattr(document, "editor_render_context_json", None), None)
    files_count = 0
    versions_count = 0
    try:
        files_count_result = await db.execute(
            select(func.count(OutgoingDocumentFile.id)).where(OutgoingDocumentFile.document_id == str(document.id))
        )
        files_count = int(files_count_result.scalar() or 0)
        versions_count_result = await db.execute(
            select(func.count(OutgoingDocumentVersion.id)).where(OutgoingDocumentVersion.document_id == str(document.id))
        )
        versions_count = int(versions_count_result.scalar() or 0)
    except Exception:
        files_count = 0
        versions_count = 0
    base = OutgoingDocumentResponse.model_validate(document).model_copy(
        update={
            "recipient_company_name": recipient.name if recipient else None,
            "deal_title": deal.title if deal else None,
            "contract_number": contract.contract_number if contract else None,
            "outgoing_number": display_number or document.outgoing_number,
            "outgoing_number_display": display_number or document.outgoing_number,
            "our_company_key": company_key,
            "document_kind": document_kind,
            "document_kind_label": DOCUMENT_KIND_LABELS.get(document_kind),
            "bank_account_snapshot": selected_bank_account,
            "selected_bank_account": selected_bank_account,
            "bank_accounts_count": len(bank_accounts),
            "linked_stage_ids": linked_stage_ids,
            "linked_payment_items": linked_payment_items,
            "editor_mode": normalize_editor_mode(getattr(document, "editor_mode", None)),
            "editor_schema_version": int(getattr(document, "editor_schema_version", None) or 1),
            "editor_draft": editor_draft,
            "editor_validation": editor_validation,
            "editor_render_context": editor_render_context,
            "files_count": files_count,
            "versions_count": versions_count,
        }
    )
    if not include_details:
        return base
    versions = await OutgoingDocumentVersion.get_by_document(db, str(document.id))
    files = await OutgoingDocumentFile.get_by_document(db, str(document.id))
    return OutgoingDocumentDetailResponse(
        **base.model_dump(),
        versions=[OutgoingDocumentVersionResponse.model_validate(version) for version in versions],
        files=[OutgoingDocumentFileResponse.model_validate(item) for item in files],
    )


async def _get_accessible_deal_for_outgoing(
    deal_id: str,
    request: Request,
    db: AsyncSession,
    user,
) -> Deal:
    deal = await Deal.get_by_id(db, deal_id)
    if not deal:
        raise HTTPException(status_code=404, detail="Deal not found")
    read_all, read_assigned = await get_section_permissions(db, user.role_id, "outgoing_registry")
    if not read_all:
        if not read_assigned:
            raise HTTPException(status_code=404, detail="Deal not found")
        allowed = await allowed_deal_ids(db, request, user)
        if allowed is not None and str(deal.id) not in set(allowed):
            raise HTTPException(status_code=404, detail="Deal not found")
    return deal


async def _pick_primary_contract_for_intro(
    db: AsyncSession,
    deal: Deal,
    our_company_key: Optional[str] = None,
) -> Optional[Contract]:
    contracts = await Contract.get_by_deal_id(db, str(deal.id))
    if not contracts:
        return None

    company_cache = {}

    async def load_company(company_id):
        if not company_id:
            return None
        cache_key = str(company_id)
        if cache_key not in company_cache:
            company_cache[cache_key] = await Company.get_by_id(db, cache_key)
        return company_cache[cache_key]

    scored_contracts = []
    for contract in contracts:
        if not contract:
            continue
        score = 0
        if deal.customer_id and contract.customer_id and str(contract.customer_id) == str(deal.customer_id):
            score += 100
        if contract.contract_type == "general_contractor":
            score += 60
        elif contract.contract_type == "services":
            score += 35
        elif contract.contract_type == "partial_contractor":
            score += 25
        elif contract.contract_type == "supply_out":
            score += 15
        if contract.status == "completed":
            score += 5
        executor_company = await load_company(contract.executor_id)
        if _company_matches_key(executor_company, our_company_key):
            score += 50
        scored_contracts.append((score, contract))

    if not scored_contracts:
        return None

    scored_contracts.sort(
        key=lambda item: (
            -item[0],
            item[1].contract_date or date.max,
            str(item[1].id),
        )
    )
    return scored_contracts[0][1]


@router.get("/", response_model=List[OutgoingDocumentResponse])
async def get_outgoing_documents(
    request: Request,
    skip: int = 0,
    limit: int = 20,
    recipient_company_id: Optional[str] = Query(None),
    deal_id: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    document_kind: Optional[str] = Query(None),
    our_company_key: Optional[str] = Query(None),
    date_from: Optional[str] = Query(None),
    date_to: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    query = select(OutgoingDocument)
    filters = []
    if recipient_company_id:
        filters.append(_id_conditions(OutgoingDocument.recipient_company_id, recipient_company_id))
    if deal_id:
        filters.append(_id_conditions(OutgoingDocument.deal_id, deal_id))
    if search:
        term = f"%{search.strip()}%"
        filters.append(
            or_(
                OutgoingDocument.subject.ilike(term),
                OutgoingDocument.outgoing_number.ilike(term),
                OutgoingDocument.attachments_list.ilike(term),
            )
        )
    if status:
        statuses = [s.strip() for s in status.split(",") if s.strip()]
        if statuses:
            filters.append(OutgoingDocument.status.in_(statuses))
    if document_kind:
        kinds = [_normalize_document_kind(item) for item in document_kind.split(",") if item.strip()]
        if kinds:
            filters.append(OutgoingDocument.document_kind.in_(kinds))
    if our_company_key:
        filters.append(OutgoingDocument.our_company_key == _normalize_company_key(our_company_key))
    if date_from:
        filters.append(OutgoingDocument.letter_date >= _parse_filter_date(date_from, "date_from"))
    if date_to:
        filters.append(OutgoingDocument.letter_date <= _parse_filter_date(date_to, "date_to"))
    if filters:
        query = query.where(and_(*filters))
    read_all, read_assigned = await get_section_permissions(db, user.role_id, "outgoing_registry")
    if not read_all:
        if not read_assigned:
            return []
        allowed = await allowed_deal_ids(db, request, user)
        if allowed == []:
            return []
        query = query.where(OutgoingDocument.deal_id.in_(allowed))
    query = query.order_by(
        func.coalesce(OutgoingDocument.outgoing_number_company_seq, 0).desc(),
        OutgoingDocument.letter_date.desc(),
        OutgoingDocument.created_at.desc(),
    ).offset(skip).limit(limit)
    result = await db.execute(query)
    documents = result.scalars().all()
    return [await _serialize_document(db, doc, include_details=False) for doc in documents]


@router.get("/deal/{deal_id}/closing-documents")
async def get_deal_closing_documents(
    deal_id: str,
    request: Request,
    exclude_document_id: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    await _get_accessible_deal_for_outgoing(deal_id, request, db, user)
    result = await db.execute(
        select(OutgoingDocument)
        .where(
            and_(
                _id_conditions(OutgoingDocument.deal_id, deal_id),
                OutgoingDocument.document_kind.in_(("invoice", "upd", "act", "vat_invoice")),
            )
        )
        .order_by(OutgoingDocument.letter_date.desc(), OutgoingDocument.created_at.desc())
    )
    documents = result.scalars().all()
    by_stage = {}
    payment_usage = {}
    serialized_documents = []

    def append_stage_document(stage_id, payload):
        normalized_stage_id = str(stage_id or "").strip()
        if not normalized_stage_id:
            return
        stage_docs = by_stage.setdefault(normalized_stage_id, [])
        for existing in stage_docs:
            if str(existing.get("id")) == str(payload.get("id")):
                if existing.get("amount") is None and payload.get("amount") is not None:
                    existing["amount"] = payload.get("amount")
                return
        stage_docs.append(payload)

    for document in documents:
        if exclude_document_id and str(document.id) == str(exclude_document_id):
            continue
        linked_stage_ids = _json_loads_or_default(getattr(document, "linked_stage_ids", None), []) or []
        linked_payment_items = _json_loads_or_default(getattr(document, "linked_payment_items", None), []) or []
        payment_total = sum(_payment_item_amount(item) for item in linked_payment_items)
        base_payload = await _closing_document_payload(
            db,
            document,
            amount=payment_total if payment_total > 0 else None,
        )
        serialized_documents.append(
            {
                **base_payload,
                "linked_stage_ids": linked_stage_ids,
                "linked_payment_items": linked_payment_items,
            }
        )
        for raw_stage_id in linked_stage_ids:
            append_stage_document(raw_stage_id, base_payload)
        for item in linked_payment_items:
            if not isinstance(item, dict):
                continue
            entry_id = str(item.get("entry_id") or item.get("payment_id") or "").strip()
            amount = _payment_item_amount(item)
            if not entry_id or amount <= 0:
                continue
            usage_row = payment_usage.setdefault(entry_id, {"used_amount": 0.0, "documents": []})
            usage_row["used_amount"] = round(float(usage_row["used_amount"]) + amount, 2)
            payment_payload = await _closing_document_payload(db, document, amount=amount)
            usage_row["documents"].append(payment_payload)
            entry_result = await db.execute(select(IncomeExpenseEntry).where(_id_conditions(IncomeExpenseEntry.id, entry_id)))
            entry = entry_result.scalar_one_or_none()
            if entry and getattr(entry, "stage_id", None):
                append_stage_document(entry.stage_id, payment_payload)

    return {
        "deal_id": str(deal_id),
        "by_stage": by_stage,
        "payment_usage": payment_usage,
        "documents": serialized_documents,
    }


@router.get("/sequences")
async def get_outgoing_sequences(
    request: Request,
    document_kind: Optional[str] = Query("letter"),
    sequence_date: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    _require_admin(request)
    normalized_kind = _normalize_document_kind(document_kind)
    parsed_date = _parse_date(sequence_date) if sequence_date else date.today()
    payload = []
    if normalized_kind == "letter":
        result = await db.execute(select(OutgoingNumberSequence))
        rows = result.scalars().all()
        for key, meta in COMPANY_OPTIONS.items():
            row = next((item for item in rows if item.our_company_key == key), None)
            payload.append(
                {
                    "our_company_key": key,
                    "label": meta["label"],
                    "document_kind": normalized_kind,
                    "sequence_date": None,
                    "next_seq": int(row.next_seq) if row else settings.OUTGOING_NUMBER_START,
                    "display_number": f"{int(row.next_seq) if row else settings.OUTGOING_NUMBER_START}/{parsed_date:%Y-%m}",
                }
            )
        return payload

    if normalized_kind not in FINANCIAL_DAILY_NUMBER_KINDS:
        return []

    for key, meta in COMPANY_OPTIONS.items():
        next_seq = await _peek_daily_financial_sequence(
            db,
            company_key=key,
            document_kind=normalized_kind,
            document_date=parsed_date,
        )
        payload.append(
            {
                "our_company_key": key,
                "label": meta["label"],
                "document_kind": normalized_kind,
                "sequence_date": parsed_date.isoformat(),
                "next_seq": int(next_seq),
                "display_number": _format_financial_number(int(next_seq), parsed_date),
            }
        )
    return payload


@router.put("/sequences/{company_key}")
async def update_outgoing_sequence(
    company_key: str,
    next_seq: int = Body(..., embed=True),
    document_kind: Optional[str] = Body("letter", embed=True),
    sequence_date: Optional[str] = Body(None, embed=True),
    request: Request = None,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    _require_admin(request)
    if next_seq < 1:
        raise HTTPException(status_code=400, detail="next_seq must be >= 1")
    normalized = _normalize_company_key(company_key)
    normalized_kind = _normalize_document_kind(document_kind)
    if normalized_kind in FINANCIAL_DAILY_NUMBER_KINDS:
        parsed_date = _parse_date(sequence_date)
        result = await db.execute(
            select(OutgoingDailyNumberSequence).where(
                and_(
                    OutgoingDailyNumberSequence.our_company_key == normalized,
                    OutgoingDailyNumberSequence.document_kind == normalized_kind,
                    OutgoingDailyNumberSequence.sequence_date == parsed_date,
                )
            )
        )
        row = result.scalar_one_or_none()
        if not row:
            row = OutgoingDailyNumberSequence(
                id=str(uuid.uuid4()),
                our_company_key=normalized,
                document_kind=normalized_kind,
                sequence_date=parsed_date,
                next_seq=next_seq,
            )
            db.add(row)
        else:
            row.next_seq = next_seq
        try:
            await db.commit()
        except OperationalError:
            await db.rollback()
            raise HTTPException(status_code=409, detail="Sequence update busy, retry")
        return {
            "our_company_key": normalized,
            "label": _company_label(normalized),
            "document_kind": normalized_kind,
            "sequence_date": parsed_date.isoformat(),
            "next_seq": int(row.next_seq),
            "display_number": _format_financial_number(int(row.next_seq), parsed_date),
        }

    if normalized_kind != "letter":
        raise HTTPException(status_code=400, detail="Selected document kind does not use manual global sequence")

    result = await db.execute(
        select(OutgoingNumberSequence).where(OutgoingNumberSequence.our_company_key == normalized)
    )
    row = result.scalar_one_or_none()
    result.close()
    if not row:
        row = OutgoingNumberSequence(our_company_key=normalized, next_seq=next_seq)
        db.add(row)
    else:
        row.next_seq = next_seq
    try:
        await db.commit()
    except OperationalError:
        await db.rollback()
        raise HTTPException(status_code=409, detail="Sequence update busy, retry")
    return {
        "our_company_key": normalized,
        "label": _company_label(normalized),
        "document_kind": normalized_kind,
        "sequence_date": None,
        "next_seq": int(row.next_seq),
    }


@router.get("/deal-intro/{deal_id}")
async def get_outgoing_deal_intro(
    deal_id: str,
    request: Request,
    recipient_short_name: Optional[str] = Query(None),
    recipient_company_id: Optional[str] = Query(None),
    our_company_key: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    deal = await _get_accessible_deal_for_outgoing(deal_id, request, db, user)

    recipient_label = (recipient_short_name or "").strip()
    if not recipient_label and recipient_company_id:
        recipient = await Company.get_by_id(db, recipient_company_id)
        if recipient:
            recipient_label = (recipient.short_name or recipient.name or "").strip()
    if not recipient_label:
        raise HTTPException(status_code=400, detail="Recipient short name is required")

    contract = await _pick_primary_contract_for_intro(db, deal, our_company_key)
    if not contract:
        raise HTTPException(status_code=404, detail="Primary deal contract not found")

    contract_number = _contract_number_with_prefix(contract.contract_number)
    contract_date = _format_ru_date(contract.contract_date)
    object_name = (deal.obj_name or deal.title or "").strip()
    if not object_name:
        raise HTTPException(status_code=400, detail="Deal object name is empty")

    company_key = _normalize_company_key(our_company_key)
    our_company_name = FORMAL_COMPANY_NAMES.get(company_key, FORMAL_COMPANY_NAMES["normbud"])
    paragraph_text = (
        f"Между {our_company_name} и {recipient_label} заключен Договор "
        f"{contract_number} от {contract_date} на выполнение работ по объекту "
        f"«{object_name}»."
    )

    return {
        "paragraph_text": paragraph_text,
        "our_company_name": our_company_name,
        "recipient_short_name": recipient_label,
        "contract_id": str(contract.id),
        "contract_number": contract.contract_number,
        "contract_date": contract.contract_date.isoformat() if contract.contract_date else None,
        "object_name": object_name,
    }


async def _get_accessible_outgoing_document(
    document_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await OutgoingDocument.get_by_id(db, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    read_all, read_assigned = await get_section_permissions(db, user.role_id, "outgoing_registry")
    if not read_all:
        if not read_assigned:
            raise HTTPException(status_code=404, detail="Document not found")
        allowed = await allowed_deal_ids(db, request, user)
        if allowed is not None:
            if not document.deal_id or str(document.deal_id) not in set(allowed):
                raise HTTPException(status_code=404, detail="Document not found")
    return document


def _build_editor_context_response(
    *,
    document_id: Optional[str],
    document_kind: str,
    editor_mode: str,
    editor_schema_version: int,
    editor_draft,
    editor_validation,
    render_payload: dict,
):
    normalized_draft = normalize_editor_draft(editor_draft, document_kind)
    normalized_validation = editor_validation or validate_editor_draft(normalized_draft, document_kind)
    return {
        "document_id": document_id,
        "document_kind": document_kind,
        "schema_version": int(editor_schema_version or 1),
        "editor_mode": normalize_editor_mode(editor_mode),
        "draft": normalized_draft,
        "validation": normalized_validation,
        "block_catalog": get_editor_block_catalog(document_kind),
        "fields": get_template_fields(module="outgoing_registry", document_kind=document_kind),
        "template_blocks": _get_structured_template_blocks(document_kind),
        "resolved_fields": _build_editor_resolved_fields(render_payload, document_kind),
        "resolved_context": render_payload,
    }


def _build_preview_outgoing_number(document_kind: str, document_date: date) -> str:
    normalized_kind = _normalize_document_kind(document_kind)
    if normalized_kind == "act":
        return "0"
    if normalized_kind == "letter":
        return _format_outgoing_number(0, document_date)
    return _format_financial_number(0, document_date)


async def _build_transient_editor_document(
    payload: OutgoingDocumentResolveRequest,
    request: Request,
    db: AsyncSession,
    user,
):
    base_document = None
    if payload.document_id:
        base_document = await _get_accessible_outgoing_document(str(payload.document_id), request, db, user)

    document_kind = _normalize_document_kind(
        payload.document_kind or getattr(base_document, "document_kind", None) or "letter"
    )
    editor_mode = normalize_editor_mode(payload.editor_mode or getattr(base_document, "editor_mode", None))
    editor_schema_version = int(
        payload.editor_schema_version
        or getattr(base_document, "editor_schema_version", None)
        or EDITOR_SCHEMA_VERSION
    )
    editor_payload = _normalize_editor_payload(
        document_kind=document_kind,
        editor_mode=editor_mode,
        editor_schema_version=editor_schema_version,
        editor_draft=payload.editor_draft if payload.editor_draft is not None else getattr(base_document, "editor_draft_json", None),
        editor_validation=payload.editor_validation if payload.editor_validation is not None else _json_loads_or_default(getattr(base_document, "editor_validation_json", None), None),
        editor_render_context=payload.editor_render_context if payload.editor_render_context is not None else _json_loads_or_default(getattr(base_document, "editor_render_context_json", None), None),
    )

    recipient_company_id = (
        str(payload.recipient_company_id)
        if payload.recipient_company_id
        else (str(base_document.recipient_company_id) if base_document and getattr(base_document, "recipient_company_id", None) else None)
    )
    if not recipient_company_id:
        return {
            "base_document": base_document,
            "document_kind": document_kind,
            "editor_mode": editor_mode,
            "editor_schema_version": editor_schema_version,
            "editor_payload": editor_payload,
            "transient_document": None,
        }

    recipient = await Company.get_by_id(db, recipient_company_id)
    if not recipient:
        raise HTTPException(status_code=404, detail="Recipient company not found")

    target_deal_id = (
        str(payload.deal_id)
        if payload.deal_id
        else (str(base_document.deal_id) if base_document and getattr(base_document, "deal_id", None) else None)
    )
    if target_deal_id:
        await _get_accessible_deal_for_outgoing(target_deal_id, request, db, user)

    target_contract_id = (
        str(payload.contract_id)
        if payload.contract_id
        else (str(base_document.contract_id) if base_document and getattr(base_document, "contract_id", None) else None)
    )
    contract = await _validate_contract_for_document(db, target_contract_id, target_deal_id)

    resolved_company_key = _normalize_company_key(
        getattr(base_document, "our_company_key", None)
        if base_document and getattr(base_document, "our_company_key", None)
        else payload.our_company_key
    )
    resolved_date = payload.letter_date or getattr(base_document, "letter_date", None) or date.today()
    bank_index, bank_snapshot = _resolve_bank_account_snapshot(
        recipient,
        payload.bank_account_index if payload.bank_account_index is not None else getattr(base_document, "bank_account_index", None),
    )
    normalized_stage_ids = await _validate_stage_links(
        db,
        payload.linked_stage_ids if payload.linked_stage_ids is not None else getattr(base_document, "linked_stage_ids", None),
        target_deal_id,
    )
    normalized_payment_items = await _validate_payment_links(
        db,
        payload.linked_payment_items if payload.linked_payment_items is not None else getattr(base_document, "linked_payment_items", None),
        target_deal_id,
        str(contract.id) if contract else target_contract_id,
        exclude_document_id=str(base_document.id) if base_document else None,
    )
    resolved_subject = (
        payload.subject
        if payload.subject is not None
        else getattr(base_document, "subject", None)
    ) or DOCUMENT_KIND_DEFAULT_SUBJECTS.get(document_kind, "Исходящий документ")
    preview_number = (
        getattr(base_document, "outgoing_number", None)
        or _build_preview_outgoing_number(document_kind, resolved_date)
    )

    transient_document = OutgoingDocument(
        id=str(base_document.id) if base_document else str(uuid.uuid4()),
        outgoing_number=preview_number,
        outgoing_number_seq=getattr(base_document, "outgoing_number_seq", 0) or 0,
        document_kind=document_kind,
        our_company_key=resolved_company_key,
        outgoing_number_company_seq=getattr(base_document, "outgoing_number_company_seq", None),
        recipient_company_id=recipient_company_id,
        deal_id=target_deal_id,
        contract_id=str(contract.id) if contract else target_contract_id,
        letter_date=resolved_date,
        subject=resolved_subject,
        body=payload.body if payload.body is not None else getattr(base_document, "body", "") or "",
        attachments_list=payload.attachments_list if payload.attachments_list is not None else getattr(base_document, "attachments_list", "") or "",
        bank_account_index=bank_index,
        bank_account_snapshot=_json_dumps_or_none(bank_snapshot),
        linked_stage_ids=_json_dumps_or_none(normalized_stage_ids),
        linked_payment_items=_json_dumps_or_none(normalized_payment_items),
        recipient_short_name=payload.recipient_short_name if payload.recipient_short_name is not None else getattr(base_document, "recipient_short_name", "") or "",
        recipient_to_name=payload.recipient_to_name if payload.recipient_to_name is not None else getattr(base_document, "recipient_to_name", "") or "",
        recipient_appeal=payload.recipient_appeal if payload.recipient_appeal is not None else getattr(base_document, "recipient_appeal", "") or "",
        recipient_eio=payload.recipient_eio if payload.recipient_eio is not None else getattr(base_document, "recipient_eio", "") or "",
        recipient_genitive_name=payload.recipient_genitive_name if payload.recipient_genitive_name is not None else getattr(base_document, "recipient_genitive_name", "") or "",
        recipient_salutation=payload.recipient_salutation if payload.recipient_salutation is not None else getattr(base_document, "recipient_salutation", "") or "",
        editor_mode=editor_payload["editor_mode"],
        editor_schema_version=editor_payload["editor_schema_version"],
        editor_draft_json=editor_payload["editor_draft_json"],
        editor_validation_json=editor_payload["editor_validation_json"],
        editor_render_context_json=editor_payload["editor_render_context_json"],
        status=payload.status or getattr(base_document, "status", None) or "draft",
    )
    return {
        "base_document": base_document,
        "document_kind": document_kind,
        "editor_mode": editor_mode,
        "editor_schema_version": editor_schema_version,
        "editor_payload": editor_payload,
        "transient_document": transient_document,
    }


@router.get("/editor/schema")
async def get_outgoing_editor_schema(
    document_kind: Optional[str] = Query("letter"),
):
    normalized_kind = _normalize_document_kind(document_kind)
    default_draft = default_editor_draft(normalized_kind)
    return {
        "document_kind": normalized_kind,
        "schema_version": EDITOR_SCHEMA_VERSION,
        "default_draft": default_draft,
        "block_catalog": get_editor_block_catalog(normalized_kind),
        "fields": get_template_fields(module="outgoing_registry", document_kind=normalized_kind),
        "template_blocks": _get_structured_template_blocks(normalized_kind),
    }


@router.get("/templates/v2")
async def list_template_v2(
    document_kind: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    """List template_v2 templates (with layout_html populated) for the outgoing registry."""
    from app.models import DocumentTemplate
    from sqlalchemy import select as sa_select, and_

    conditions = [
        DocumentTemplate.module == "outgoing_registry",
        DocumentTemplate.layout_html.isnot(None),
        DocumentTemplate.is_active == True,  # noqa: E712
    ]
    if document_kind:
        conditions.append(DocumentTemplate.document_kind == _normalize_document_kind(document_kind))

    rows = (await db.execute(sa_select(DocumentTemplate).where(and_(*conditions)))).scalars().all()
    return [
        {
            "id": t.id,
            "name": t.name,
            "description": t.description,
            "module": t.module,
            "document_kind": t.document_kind,
            "our_company_key": t.our_company_key,
            "layout_html": t.layout_html,
            "editable_regions": t.editable_regions_json or [],
            "placeholder_fields": t.placeholder_fields_json or [],
        }
        for t in rows
    ]


@router.get("/templates/v2/{template_id}")
async def get_template_v2(
    template_id: str,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    """Get a single template_v2 by id with full layout/fields payload."""
    from app.models import DocumentTemplate
    from sqlalchemy import select as sa_select

    t = (await db.execute(sa_select(DocumentTemplate).where(DocumentTemplate.id == template_id))).scalar_one_or_none()
    if t is None or not t.layout_html:
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    return {
        "id": t.id,
        "name": t.name,
        "description": t.description,
        "module": t.module,
        "document_kind": t.document_kind,
        "our_company_key": t.our_company_key,
        "layout_html": t.layout_html,
        "editable_regions": t.editable_regions_json or [],
        "placeholder_fields": t.placeholder_fields_json or [],
    }


def _serialize_template_v2(t) -> dict:
    return {
        "id": t.id,
        "name": t.name,
        "description": t.description,
        "module": t.module,
        "document_kind": t.document_kind,
        "our_company_key": t.our_company_key,
        "layout_html": t.layout_html,
        "editable_regions": t.editable_regions_json or [],
        "placeholder_fields": t.placeholder_fields_json or [],
    }


_DEFAULT_EDITABLE_V2 = [{
    "key": "body",
    "label": "Текст документа",
    "allowed_marks": ["bold", "italic", "underline", "list", "alignment", "link"],
    "default_html": "<p></p>",
}]


@router.post("/templates/v2")
async def create_template_v2(
    payload: dict,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    """Create a new template_v2 (layout_html-driven)."""
    from app.models import DocumentTemplate

    name = (payload.get("name") or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="Укажите название шаблона")
    document_kind = _normalize_document_kind(payload.get("document_kind") or "letter")
    fields = payload.get("placeholder_fields") or []
    t = DocumentTemplate(
        name=name,
        description=payload.get("description") or f"template_v2 — {document_kind}",
        module="outgoing_registry",
        document_kind=document_kind,
        our_company_key=payload.get("our_company_key") or None,
        binding_type="global",
        status="approved",
        is_active=True,
        fields_json=[f.get("key") for f in fields if isinstance(f, dict) and f.get("key")],
        layout_html=payload.get("layout_html") or "<p data-editable=\"body\"></p>",
        editable_regions_json=payload.get("editable_regions") or _DEFAULT_EDITABLE_V2,
        placeholder_fields_json=fields,
        created_by=str(user.id),
    )
    db.add(t)
    await db.commit()
    await db.refresh(t)
    return _serialize_template_v2(t)


@router.put("/templates/v2/{template_id}")
async def update_template_v2(
    template_id: str,
    payload: dict,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    """Update a template_v2 layout / fields / meta."""
    from app.models import DocumentTemplate
    from sqlalchemy import select as sa_select

    t = (await db.execute(sa_select(DocumentTemplate).where(DocumentTemplate.id == template_id))).scalar_one_or_none()
    if t is None:
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    if "name" in payload and str(payload["name"]).strip():
        t.name = str(payload["name"]).strip()
    if "description" in payload:
        t.description = payload["description"]
    if "document_kind" in payload and payload["document_kind"]:
        t.document_kind = _normalize_document_kind(payload["document_kind"])
    if "our_company_key" in payload:
        t.our_company_key = payload["our_company_key"] or None
    if "layout_html" in payload:
        t.layout_html = payload["layout_html"] or ""
    if "editable_regions" in payload:
        t.editable_regions_json = payload["editable_regions"] or _DEFAULT_EDITABLE_V2
    if "placeholder_fields" in payload:
        fields = payload["placeholder_fields"] or []
        t.placeholder_fields_json = fields
        t.fields_json = [f.get("key") for f in fields if isinstance(f, dict) and f.get("key")]
    t.updated_by = str(user.id)
    await db.commit()
    await db.refresh(t)
    return _serialize_template_v2(t)


@router.delete("/templates/v2/{template_id}")
async def delete_template_v2(
    template_id: str,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    from app.models import DocumentTemplate
    from sqlalchemy import select as sa_select

    t = (await db.execute(sa_select(DocumentTemplate).where(DocumentTemplate.id == template_id))).scalar_one_or_none()
    if t is None:
        raise HTTPException(status_code=404, detail="Шаблон не найден")
    await db.delete(t)
    await db.commit()
    return {"ok": True}


@router.post("/editor/resolve")
async def resolve_outgoing_editor_context(
    payload: OutgoingDocumentResolveRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    preview_context = await _build_transient_editor_document(payload, request, db, user)
    base_document = preview_context["base_document"]
    document_kind = preview_context["document_kind"]
    editor_mode = preview_context["editor_mode"]
    editor_schema_version = preview_context["editor_schema_version"]
    editor_payload = preview_context["editor_payload"]
    transient_document = preview_context["transient_document"]

    if transient_document is None:
        return _build_editor_context_response(
            document_id=str(base_document.id) if base_document else None,
            document_kind=document_kind,
            editor_mode=editor_mode,
            editor_schema_version=editor_schema_version,
            editor_draft=_json_loads_or_default(editor_payload["editor_draft_json"], None),
            editor_validation=_json_loads_or_default(editor_payload["editor_validation_json"], None),
            render_payload={},
        )
    render_payload = await _build_document_render_payload(db, transient_document)
    return _build_editor_context_response(
        document_id=str(base_document.id) if base_document else None,
        document_kind=document_kind,
        editor_mode=editor_mode,
        editor_schema_version=editor_schema_version,
        editor_draft=_json_loads_or_default(editor_payload["editor_draft_json"], None),
        editor_validation=_json_loads_or_default(editor_payload["editor_validation_json"], None),
        render_payload=render_payload,
    )


@router.post("/editor/preview-pdf")
async def preview_outgoing_editor_pdf(
    payload: OutgoingDocumentResolveRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    preview_context = await _build_transient_editor_document(payload, request, db, user)
    transient_document = preview_context["transient_document"]
    if transient_document is None:
        raise HTTPException(status_code=400, detail="Recipient company is required for preview")
    _, pdf_bytes = await _render_document_files(db, transient_document)
    if pdf_bytes is None:
        raise HTTPException(status_code=409, detail="Preview PDF generation failed")
    filename = clean_name(f"{_build_outgoing_file_base_clean(transient_document)}_preview.pdf")
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": _content_disposition("inline", filename),
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@router.post("/editor/preview-docx")
async def preview_outgoing_editor_docx(
    payload: OutgoingDocumentResolveRequest,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    preview_context = await _build_transient_editor_document(payload, request, db, user)
    transient_document = preview_context["transient_document"]
    if transient_document is None:
        raise HTTPException(status_code=400, detail="Recipient company is required for preview")
    docx_bytes, _ = await _render_document_files(db, transient_document)
    if docx_bytes is None:
        raise HTTPException(status_code=409, detail="Preview DOCX generation failed")
    filename = clean_name(f"{_build_outgoing_file_base_clean(transient_document)}_preview.docx")
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": _content_disposition("inline", filename),
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@router.get("/{document_id}/editor-context")
async def get_outgoing_editor_context(
    document_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await _get_accessible_outgoing_document(document_id, request, db, user)
    document_kind = _normalize_document_kind(getattr(document, "document_kind", None))
    render_payload = await _build_document_render_payload(db, document)
    editor_draft = _json_loads_or_default(getattr(document, "editor_draft_json", None), None)
    editor_validation = _json_loads_or_default(getattr(document, "editor_validation_json", None), None)
    return _build_editor_context_response(
        document_id=str(document.id),
        document_kind=document_kind,
        editor_mode=getattr(document, "editor_mode", None) or "classic",
        editor_schema_version=int(getattr(document, "editor_schema_version", None) or 1),
        editor_draft=editor_draft,
        editor_validation=editor_validation,
        render_payload=render_payload,
    )


@router.get("/{document_id}", response_model=OutgoingDocumentDetailResponse)
async def get_outgoing_document(
    document_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await _get_accessible_outgoing_document(document_id, request, db, user)
    return await _serialize_document(db, document, include_details=True)


@router.get("/{document_id}/preview-pdf")
async def preview_outgoing_document_pdf(
    document_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await _get_accessible_outgoing_document(document_id, request, db, user)
    effective_render = await _resolve_effective_render_payload(db, document)
    pdf_bytes = effective_render["pdf_bytes"]
    filename = clean_name(f"{_build_outgoing_file_base_clean(document)}_preview.pdf")
    if effective_render["pdf_filename"]:
        filename = effective_render["pdf_filename"] or filename
    if pdf_bytes is None:
        raise HTTPException(
            status_code=409,
            detail="Current render cache is missing. Save the document to regenerate preview.",
        )
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": _content_disposition("inline", filename),
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@router.post("/{document_id}/preview-uploaded-docx-pdf")
async def preview_uploaded_outgoing_docx_pdf(
    document_id: str,
    request: Request,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await _get_accessible_outgoing_document(document_id, request, db, user)
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Expected .docx file")
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="DOCX file is empty")
    pdf_bytes = await run_in_threadpool(_convert_docx_to_pdf_bytes, data)
    if not pdf_bytes:
        raise HTTPException(status_code=500, detail="Не удалось сформировать PDF из Word-файла")
    filename = clean_name(f"{_build_outgoing_file_base_clean(document)}_preview.pdf")
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={
            "Content-Disposition": _content_disposition("inline", filename),
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@router.post("/{document_id}/render-cache", response_model=OutgoingDocumentDetailResponse)
async def update_outgoing_document_render_cache(
    document_id: str,
    request: Request,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await _get_accessible_outgoing_document(document_id, request, db, user)
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Expected .docx file")
    docx_bytes = await file.read()
    if not docx_bytes:
        raise HTTPException(status_code=400, detail="DOCX file is empty")
    pdf_bytes = await run_in_threadpool(_convert_docx_to_pdf_bytes, docx_bytes)
    if not pdf_bytes:
        raise HTTPException(status_code=500, detail="Не удалось сформировать PDF из Word-файла")
    return await _store_current_render_files(db, document, docx_bytes, pdf_bytes)


@router.post("/{document_id}/render-cache/generate", response_model=OutgoingDocumentDetailResponse)
async def generate_outgoing_document_render_cache(
    document_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await _get_accessible_outgoing_document(document_id, request, db, user)
    docx_bytes, pdf_bytes = await _render_document_files(db, document)
    if not docx_bytes or not pdf_bytes:
        raise HTTPException(status_code=500, detail="Не удалось сформировать файлы документа")
    return await _store_current_render_files(db, document, docx_bytes, pdf_bytes)


@router.get("/{document_id}/render-docx")
async def render_outgoing_document_docx(
    document_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await _get_accessible_outgoing_document(document_id, request, db, user)
    effective_render = await _resolve_effective_render_payload(db, document)
    docx_bytes = effective_render["docx_bytes"]
    filename = clean_name(f"{_build_outgoing_file_base_clean(document)}.docx")
    if effective_render["docx_filename"]:
        filename = effective_render["docx_filename"] or filename
    if docx_bytes is None:
        raise HTTPException(
            status_code=409,
            detail="Current render cache is missing. Save the document to regenerate Word file.",
        )
    if not docx_bytes:
        raise HTTPException(status_code=500, detail="Не удалось сформировать Word-файл")
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": _content_disposition("attachment", filename),
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
            "Expires": "0",
        },
    )


@router.post("/", response_model=OutgoingDocumentDetailResponse)
async def create_outgoing_document(
    document_kind: Optional[str] = Form("letter"),
    editor_mode: Optional[str] = Form("classic"),
    editor_schema_version: Optional[int] = Form(1),
    editor_draft: Optional[str] = Form(None),
    editor_validation: Optional[str] = Form(None),
    editor_render_context: Optional[str] = Form(None),
    recipient_company_id: str = Form(...),
    deal_id: Optional[str] = Form(None),
    contract_id: Optional[str] = Form(None),
    letter_date: Optional[str] = Form(None),
    subject: Optional[str] = Form(None),
    body: Optional[str] = Form(""),
    attachments_list: Optional[str] = Form(""),
    bank_account_index: Optional[int] = Form(None),
    linked_stage_ids: Optional[str] = Form(None),
    linked_payment_items: Optional[str] = Form(None),
    recipient_short_name: Optional[str] = Form(""),
    recipient_to_name: Optional[str] = Form(""),
    recipient_appeal: Optional[str] = Form(""),
    recipient_eio: Optional[str] = Form(""),
    recipient_genitive_name: Optional[str] = Form(""),
    recipient_salutation: Optional[str] = Form(""),
    status: Optional[str] = Form("draft"),
    our_company_key: Optional[str] = Form(None),
    attachments_files: Optional[List[UploadFile]] = File(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    recipient = await Company.get_by_id(db, recipient_company_id)
    if not recipient:
        raise HTTPException(status_code=404, detail="Recipient company not found")
    if deal_id:
        deal = await Deal.get_by_id(db, deal_id)
        if not deal:
            raise HTTPException(status_code=404, detail="Deal not found")

    parsed_date = _parse_date(letter_date)
    normalized_kind = _normalize_document_kind(document_kind)
    company_key = _normalize_company_key(our_company_key)
    editor_payload = _normalize_editor_payload(
        document_kind=normalized_kind,
        editor_mode=editor_mode,
        editor_schema_version=editor_schema_version,
        editor_draft=editor_draft,
        editor_validation=_json_loads_or_default(editor_validation, None),
        editor_render_context=_json_loads_or_default(editor_render_context, None),
    )
    contract = await _validate_contract_for_document(db, contract_id, deal_id)
    bank_index, bank_snapshot = _resolve_bank_account_snapshot(recipient, bank_account_index)
    normalized_stage_ids = await _validate_stage_links(db, linked_stage_ids, deal_id)
    normalized_payment_items = await _validate_payment_links(db, linked_payment_items, deal_id, contract_id)
    resolved_subject = (subject or "").strip() or DOCUMENT_KIND_DEFAULT_SUBJECTS.get(normalized_kind, "Исходящий документ")
    act_contract_document = None
    if normalized_kind == "act":
        if not contract:
            raise HTTPException(status_code=400, detail="Act requires selected contract")
        # number_in_contract has no UNIQUE constraint, so allocate it and
        # persist the act ContractDocument atomically under the same lock
        # the contracts.py / uploads.py upload paths use — closes the
        # MAX()+1 race within and across those endpoints.
        async with sequence_lock("contract_document_number"):
            financial_seq = await _get_next_contract_act_number(db, str(contract.id))
            act_contract_document = await ContractDocument.create(
                db,
                contract_id=contract.id,
                doc_type="act",
                number_in_contract=int(financial_seq),
                status=status or "draft",
            )
        display_number = str(financial_seq)
    elif normalized_kind == "letter":
        financial_seq = await _get_company_sequence(db, company_key)
        display_number = _format_outgoing_number(financial_seq, parsed_date)
    else:
        financial_seq = await _get_daily_financial_sequence(
            db,
            company_key=company_key,
            document_kind=normalized_kind,
            document_date=parsed_date,
        )
        display_number = _format_financial_number(financial_seq, parsed_date)
    document = None
    for _ in range(3):
        seq = await _get_next_number(db)
        if normalized_kind == "letter":
            outgoing_number = f"{company_key}:{display_number}"
        else:
            outgoing_number = f"{company_key}:{normalized_kind}:{display_number}"
        try:
            document = await OutgoingDocument.create(
                db,
                outgoing_number_seq=seq,
                outgoing_number=outgoing_number,
                document_kind=normalized_kind,
                our_company_key=company_key,
                outgoing_number_company_seq=financial_seq,
                recipient_company_id=recipient_company_id,
                deal_id=deal_id,
                contract_id=str(contract.id) if contract else contract_id,
                letter_date=parsed_date,
                subject=resolved_subject,
                body=body,
                attachments_list=attachments_list,
                bank_account_index=bank_index,
                bank_account_snapshot=_json_dumps_or_none(bank_snapshot),
                linked_stage_ids=_json_dumps_or_none(normalized_stage_ids),
                linked_payment_items=_json_dumps_or_none(normalized_payment_items),
                recipient_short_name=recipient_short_name,
                recipient_to_name=recipient_to_name,
                recipient_appeal=recipient_appeal,
                recipient_eio=recipient_eio,
                recipient_genitive_name=recipient_genitive_name,
                recipient_salutation=recipient_salutation,
                editor_mode=editor_payload["editor_mode"],
                editor_schema_version=editor_payload["editor_schema_version"],
                editor_draft_json=editor_payload["editor_draft_json"],
                editor_validation_json=editor_payload["editor_validation_json"],
                editor_render_context_json=editor_payload["editor_render_context_json"],
                status=status or "draft",
            )
            break
        except IntegrityError:
            await db.rollback()
            if normalized_kind not in {"letter", "act"}:
                financial_seq = await _get_daily_financial_sequence(
                    db,
                    company_key=company_key,
                    document_kind=normalized_kind,
                    document_date=parsed_date,
                )
                display_number = _format_financial_number(financial_seq, parsed_date)
            continue
    if not document:
        # Outgoing seq couldn't be allocated — drop the pre-created act
        # ContractDocument so it doesn't orphan.
        if act_contract_document is not None:
            await ContractDocument.delete(db, str(act_contract_document.id))
        raise HTTPException(status_code=409, detail="Unable to allocate unique outgoing number, retry")

    if normalized_kind == "act" and contract and act_contract_document is not None:
        document = await OutgoingDocument.update(
            db,
            document.id,
            act_contract_document_id=str(act_contract_document.id),
        )

    if attachments_files:
        if not storage_available():
            raise HTTPException(status_code=500, detail="Storage is not configured")
        paths = _build_paths(outgoing_number)
        await ensure_path(paths["attachments"])
        for upload in attachments_files:
            content = await upload.read()
            file_name = clean_name(upload.filename)
            file_path = f"{paths['attachments']}/{file_name}"
            await upload_bytes_with_safe_extension(file_path, content)
            public_url = await publish(file_path)
            await OutgoingDocumentFile.create(
                db,
                document_id=document.id,
                file_type="attachment",
                file_path=file_path,
                file_name=file_name,
                public_url=public_url,
            )
    try:
        await _sync_document_registry(db, document)
    except Exception:
        pass

    try:
        await log_event(
            db,
            entity_type="deal" if document.deal_id else "outgoing",
            entity_id=str(document.deal_id or document.id),
            action="outgoing.create",
            created_by=str(getattr(user, "id", "")),
            details={
                "outgoing_id": str(document.id),
                "outgoing_number": document.outgoing_number,
                "deal_id": str(document.deal_id) if document.deal_id else None,
            },
        )
    except Exception:
        pass
    await safe_refresh_deal_health_issues(db, document.deal_id)
    if not document.deal_id:
        await safe_refresh_orphan_health_issues(db)

    # Event Bus v2: outbox-эмиссия для внешних подписчиков (Диадок,
    # архив документов, BI). После create — общий стандартный event.
    from app.services.event_outbox import emit_event_safe
    await emit_event_safe(
        db,
        event_type="outgoing_document.after_create",
        entity_type="outgoing_document",
        entity_id=str(document.id),
        payload={
            "id": str(document.id),
            "outgoing_number": document.outgoing_number,
            "deal_id": str(document.deal_id) if document.deal_id else None,
            "document_kind": getattr(document, "document_kind", None),
            "status": getattr(document, "status", None),
        },
    )
    return await _serialize_document(db, document, include_details=True)


@router.put("/{document_id}", response_model=OutgoingDocumentDetailResponse)
async def update_outgoing_document(
    document_id: str,
    document_update: OutgoingDocumentUpdate,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await OutgoingDocument.get_by_id(db, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    await ensure_can_edit_record(db, request, user, "outgoing_registry", document)
    previous_deal_id = str(document.deal_id) if document.deal_id else None
    if document_update.recipient_company_id:
        recipient = await Company.get_by_id(db, str(document_update.recipient_company_id))
        if not recipient:
            raise HTTPException(status_code=404, detail="Recipient company not found")
    if document_update.deal_id:
        deal = await Deal.get_by_id(db, str(document_update.deal_id))
        if not deal:
            raise HTTPException(status_code=404, detail="Deal not found")
    update_payload = document_update.dict(exclude_unset=True)
    outgoing_suffix = update_payload.pop("outgoing_number_suffix", None)
    update_payload.pop("document_kind", None)
    # Avoid writing nulls into required fields from partial updates.
    if "recipient_company_id" in update_payload and not update_payload.get("recipient_company_id"):
        update_payload.pop("recipient_company_id", None)
    if "letter_date" in update_payload and not update_payload.get("letter_date"):
        update_payload.pop("letter_date", None)
    if "deal_id" in update_payload and not update_payload.get("deal_id"):
        update_payload.pop("deal_id", None)
    if "contract_id" in update_payload and not update_payload.get("contract_id"):
        update_payload["contract_id"] = None
    target_deal_id = (
        str(update_payload["deal_id"]) if "deal_id" in update_payload and update_payload.get("deal_id")
        else (str(document.deal_id) if document.deal_id else None)
    )
    target_contract_id = (
        str(update_payload["contract_id"]) if "contract_id" in update_payload and update_payload.get("contract_id")
        else (None if "contract_id" in update_payload else (str(document.contract_id) if document.contract_id else None))
    )
    await _validate_contract_for_document(db, target_contract_id, target_deal_id)
    if update_payload.get("status") == "sent":
        await ensure_entity_action_allowed(
            db,
            entity_type="outgoing_document",
            entity_id=str(document.id),
            action_label="отправка документа",
        )
    if "bank_account_index" in update_payload or "recipient_company_id" in update_payload:
        recipient_id = str(update_payload.get("recipient_company_id") or document.recipient_company_id)
        recipient = await Company.get_by_id(db, recipient_id)
        if not recipient:
            raise HTTPException(status_code=404, detail="Recipient company not found")
        bank_index, bank_snapshot = _resolve_bank_account_snapshot(
            recipient,
            update_payload.get("bank_account_index") if "bank_account_index" in update_payload else document.bank_account_index,
        )
        update_payload["bank_account_index"] = bank_index
        update_payload["bank_account_snapshot"] = _json_dumps_or_none(bank_snapshot)
    if "linked_stage_ids" in update_payload:
        update_payload["linked_stage_ids"] = _json_dumps_or_none(
            await _validate_stage_links(db, update_payload.get("linked_stage_ids"), target_deal_id)
        )
    if "linked_payment_items" in update_payload:
        update_payload["linked_payment_items"] = _json_dumps_or_none(
            await _validate_payment_links(
                db,
                update_payload.get("linked_payment_items"),
                target_deal_id,
                target_contract_id,
                exclude_document_id=document.id,
            )
        )
    if "our_company_key" in update_payload:
        if document.our_company_key:
            update_payload.pop("our_company_key", None)
        else:
            update_payload["our_company_key"] = _normalize_company_key(update_payload.get("our_company_key"))
    existing_kind = _normalize_document_kind(getattr(document, "document_kind", None))
    if any(key in update_payload for key in ("editor_mode", "editor_schema_version", "editor_draft", "editor_validation", "editor_render_context")):
        editor_payload = _normalize_editor_payload(
            document_kind=existing_kind,
            editor_mode=update_payload.pop("editor_mode", None) or getattr(document, "editor_mode", None),
            editor_schema_version=update_payload.pop("editor_schema_version", None) or getattr(document, "editor_schema_version", None),
            editor_draft=update_payload.pop("editor_draft", None) if "editor_draft" in update_payload else getattr(document, "editor_draft_json", None),
            editor_validation=update_payload.pop("editor_validation", None) if "editor_validation" in update_payload else _json_loads_or_default(getattr(document, "editor_validation_json", None), None),
            editor_render_context=update_payload.pop("editor_render_context", None) if "editor_render_context" in update_payload else _json_loads_or_default(getattr(document, "editor_render_context_json", None), None),
        )
        update_payload.update(editor_payload)
    if outgoing_suffix and existing_kind == "letter":
        suffix_value = _parse_outgoing_suffix(outgoing_suffix)
        update_payload["outgoing_number"] = _build_outgoing_number_with_suffix(document, suffix_value)
    try:
        updated = await OutgoingDocument.update(db, document_id, **update_payload)
    except IntegrityError:
        raise HTTPException(status_code=422, detail="Invalid outgoing registry payload")
    if not updated:
        raise HTTPException(status_code=404, detail="Document not found")
    try:
        await _sync_document_registry(db, updated)
    except Exception:
        pass
    await safe_refresh_deal_health_issues(db, updated.deal_id)
    if previous_deal_id and previous_deal_id != str(updated.deal_id or ""):
        await safe_refresh_deal_health_issues(db, previous_deal_id)
    if not updated.deal_id:
        await safe_refresh_orphan_health_issues(db)

    # Event Bus v2: общее обновление документа. status-change события
    # эмитим отдельно в местах смены статуса (отдельным PR).
    from app.services.event_outbox import emit_event_safe
    await emit_event_safe(
        db,
        event_type="outgoing_document.after_update",
        entity_type="outgoing_document",
        entity_id=str(updated.id),
        payload={
            "id": str(updated.id),
            "outgoing_number": updated.outgoing_number,
            "deal_id": str(updated.deal_id) if updated.deal_id else None,
            "document_kind": getattr(updated, "document_kind", None),
            "status": getattr(updated, "status", None),
        },
    )
    return await _serialize_document(db, updated, include_details=True)


@router.delete("/{document_id}")
async def delete_outgoing_document(
    document_id: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    document = await OutgoingDocument.get_by_id(db, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    await ensure_can_edit_record(db, request, user, "outgoing_registry", document)

    if storage_available():
        try:
            paths = _build_paths(document.outgoing_number)
            await delete_path(paths["root"])
        except Exception:
            pass

    await db.execute(delete(OutgoingDocumentFile).where(OutgoingDocumentFile.document_id == document.id))
    await db.execute(delete(OutgoingDocumentVersion).where(OutgoingDocumentVersion.document_id == document.id))
    await db.execute(
        delete(Document).where(
            and_(
                Document.source_type == "outgoing_registry",
                Document.source_id == str(document.id),
            )
        )
    )
    await db.execute(delete(OutgoingDocument).where(OutgoingDocument.id == document.id))
    await db.commit()
    await safe_refresh_deal_health_issues(db, document.deal_id)
    if not document.deal_id:
        await safe_refresh_orphan_health_issues(db)

    # Event Bus v2: after_delete для уведомления внешних реестров.
    from app.services.event_outbox import emit_event_safe
    await emit_event_safe(
        db,
        event_type="outgoing_document.after_delete",
        entity_type="outgoing_document",
        entity_id=str(document.id),
        payload={
            "id": str(document.id),
            "outgoing_number": document.outgoing_number,
            "deal_id": str(document.deal_id) if document.deal_id else None,
        },
    )
    return {"message": "Document deleted"}


@router.post("/{document_id}/attachments", response_model=List[OutgoingDocumentFileResponse])
async def upload_outgoing_attachments(
    document_id: str,
    attachments_files: List[UploadFile] = File(...),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    if not storage_available():
        raise HTTPException(status_code=500, detail="Storage is not configured")
    document = await OutgoingDocument.get_by_id(db, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    if not attachments_files:
        raise HTTPException(status_code=400, detail="Files are required")

    paths = _build_paths(document.outgoing_number)
    await ensure_path(paths["attachments"])
    created = []
    for upload in attachments_files:
        content = await upload.read()
        file_name = clean_name(upload.filename)
        file_path = f"{paths['attachments']}/{file_name}"
        await upload_bytes_with_safe_extension(file_path, content)
        public_url = await publish(file_path)
        created.append(
            await OutgoingDocumentFile.create(
                db,
                document_id=document.id,
                file_type="attachment",
                file_path=file_path,
                file_name=file_name,
                public_url=public_url,
            )
        )
    try:
        await log_event(
            db,
            entity_type="deal" if document.deal_id else "outgoing",
            entity_id=str(document.deal_id or document.id),
            action="outgoing.attach",
            created_by=str(getattr(user, "id", "")),
            details={
                "outgoing_id": str(document.id),
                "outgoing_number": document.outgoing_number,
                "deal_id": str(document.deal_id) if document.deal_id else None,
            },
        )
    except Exception:
        pass
    return [OutgoingDocumentFileResponse.model_validate(item) for item in created]


@router.post("/{document_id}/versions", response_model=OutgoingDocumentVersionResponse)
async def create_outgoing_version(
    document_id: str,
    file: Optional[UploadFile] = File(None),
    comment: Optional[str] = Form(None),
    created_by: Optional[str] = Form(None),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    """
    Create a new version from uploaded DOCX file or from canonical HTML render.
    """
    document = await OutgoingDocument.get_by_id(db, document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    await ensure_entity_action_allowed(
        db,
        entity_type="outgoing_document",
        entity_id=str(document.id),
        action_label="выпуск версии документа",
    )
    previous_deal_id = str(document.deal_id) if document.deal_id else None

    # version_number has no UNIQUE constraint; serialize MAX()+1 -> the
    # version record INSERT so two concurrent "выпуск версии" calls for
    # the same document can't produce duplicate version numbers. Rare
    # manual action, so holding the lock across the render is acceptable.
    async with sequence_lock("outgoing_doc_version"):
        versions = await OutgoingDocumentVersion.get_by_document(db, document_id)
        next_version = max([v.version_number for v in versions], default=0) + 1

        pdf_bytes = None
        docx_bytes = None
        if file and file.filename.lower().endswith(".docx"):
            try:
                docx_bytes = await file.read()
                pdf_bytes = await run_in_threadpool(_convert_docx_to_pdf_bytes, docx_bytes)
            except Exception:
                pdf_bytes = None

        if pdf_bytes is None:
            effective_render = await _resolve_effective_render_payload(db, document)
            docx_bytes = effective_render["docx_bytes"]
            pdf_bytes = effective_render["pdf_bytes"]

        if pdf_bytes is None:
            raise HTTPException(
                status_code=409,
                detail="Current render cache is missing. Save the document to regenerate version files.",
            )

        # Upload PDF to storage
        if not storage_available():
            raise HTTPException(status_code=500, detail="Storage is not configured")
        paths = _build_paths(document.outgoing_number)
        await ensure_path(paths["versions"])
        base_name = _build_outgoing_file_base_clean(document)
        pdf_name = clean_name(f"{base_name}_v{next_version}.pdf")
        pdf_path = f"{paths['versions']}/{pdf_name}"
        await upload_bytes_with_safe_extension(pdf_path, pdf_bytes)
        pdf_url = await publish(pdf_path)

        docx_path = None
        docx_name = None
        docx_url = None
        if docx_bytes:
            docx_name = clean_name(f"{base_name}_v{next_version}.docx")
            docx_path = f"{paths['versions']}/{docx_name}"
            await upload_bytes_with_safe_extension(docx_path, docx_bytes)
            docx_url = await publish(docx_path)

        # Create version record
        version = await OutgoingDocumentVersion.create(
            db,
            document_id=document.id,
            version_number=next_version,
            status="draft",
            created_by=created_by,
            comment=comment,
            pdf_path=pdf_path,
            pdf_public_url=pdf_url,
        )
    if docx_path:
        await OutgoingDocumentFile.create(
            db,
            document_id=document.id,
            version_id=version.id,
            file_type="docx",
            file_path=docx_path,
            file_name=docx_name,
            public_url=docx_url,
        )
    await OutgoingDocumentFile.create(
        db,
        document_id=document.id,
        version_id=version.id,
        file_type="pdf",
        file_path=pdf_path,
        file_name=pdf_name,
        public_url=pdf_url,
    )
    try:
        await log_event(
            db,
            entity_type="deal" if document.deal_id else "outgoing",
            entity_id=str(document.deal_id or document.id),
            action="outgoing.version",
            created_by=str(getattr(user, "id", "")),
            details={
                "outgoing_id": str(document.id),
                "outgoing_number": document.outgoing_number,
                "deal_id": str(document.deal_id) if document.deal_id else None,
            },
        )
    except Exception:
        pass
    return OutgoingDocumentVersionResponse.model_validate(version)


@router.get("/{document_id}/versions/{version_id}/download")
async def download_outgoing_version(
    document_id: str,
    version_id: str,
    format: str = Query("pdf"),
    db: AsyncSession = Depends(get_db),
    user=Depends(CurrentUser),
):
    """Stream a stored version file (pdf or docx) by version id.

    Both formats are snapshotted at version creation as OutgoingDocumentFile
    rows (file_type 'pdf' / 'docx'). Streaming the bytes avoids relying on a
    public URL being externally reachable.
    """
    fmt = (format or "pdf").strip().lower()
    if fmt not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="format must be pdf or docx")

    files = await OutgoingDocumentFile.get_by_document(db, str(document_id))
    match = next(
        (f for f in files
         if str(getattr(f, "version_id", "")) == str(version_id)
         and (getattr(f, "file_type", "") or "") == fmt),
        None,
    )
    if match is None:
        raise HTTPException(
            status_code=404,
            detail=("DOCX-снимок для этой версии не сохранён" if fmt == "docx"
                    else "PDF-снимок для этой версии не найден"),
        )
    try:
        content = await read_file_bytes(match.file_path)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Файл версии не найден в хранилище")

    filename = match.file_name or f"version.{fmt}"
    media = (
        "application/pdf" if fmt == "pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    return Response(
        content,
        media_type=media,
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename, safe='')}",
            "Cache-Control": "no-store",
        },
    )
