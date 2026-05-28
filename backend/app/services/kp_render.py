"""
Серверный рендер docx-шаблона КП через python-docx.

В отличие от фронтового docxtemplater умеет вставить НАСТОЯЩУЮ docx-таблицу
на место плейсхолдера `{products_table}` — с заголовком, рамками и
оформлением, максимально близким к прилагающемуся шаблону-образцу
(Times New Roman 11pt, центр шапки, ширины колонок, итого с merge 0..3).

Поддерживает:
  • простые плейсхолдеры в формате `{snake_case}` — заменяются
    in-place в параграфах и в ячейках уже существующих таблиц;
  • специальный плейсхолдер `{products_table}` — параграф-маркер
    удаляется, на его место вставляется таблица со строками товаров.
"""
import io
import re
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PLACEHOLDER_RE = re.compile(r"\{[a-z_][a-z0-9_]*\}", re.IGNORECASE)

# Параметры таблицы — подобраны под шаблон-образец.
FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 11
# Сумма ширин ~17.1 см = A4 (21 см) минус ~2 см левый/правый margin.
COL_WIDTHS_CM = [1.0, 9.0, 1.8, 1.8, 3.5]
# Точные тексты шапки из образца. "\n" в первой ячейке даёт явный
# перенос «№ / п/п», как в шаблоне.
HEADERS = [
    "№\nп/п",
    "Наименование работ/услуг",
    "Ед. изм.",
    "Кол-во",
    "Стоимость, с НДС, руб.",
]


# ────────────────────────────────────────────────────────────────────
# Низкоуровневые помощники
# ────────────────────────────────────────────────────────────────────

def _collapse_runs(paragraph) -> None:
    """Склеить все runs параграфа в первый — иначе плейсхолдеры,
    разбитые между runs, не находятся простым str.replace."""
    runs = paragraph.runs
    if len(runs) <= 1:
        return
    first = runs[0]
    full_text = "".join(r.text or "" for r in runs)
    first.text = full_text
    for r in runs[1:]:
        r.text = ""


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    """Подставить все `{key}` плейсхолдеры в одном параграфе."""
    if "{" not in paragraph.text:
        return
    _collapse_runs(paragraph)
    if not paragraph.runs:
        return
    run = paragraph.runs[0]
    text = run.text or ""

    def repl(match):
        key = match.group(0)[1:-1]  # strip { }
        value = mapping.get(key)
        if value is None:
            return match.group(0)
        return str(value)

    new_text = PLACEHOLDER_RE.sub(repl, text)
    if new_text != text:
        run.text = new_text


def _replace_in_table(table, mapping: Dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, mapping)
            for inner in cell.tables:
                _replace_in_table(inner, mapping)


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_font(run, name: str, size_pt: int, bold: bool) -> None:
    """Назначить шрифт run-у. Кириллица нередко требует явного rFonts
    с ascii/hAnsi/cs/eastAsia — иначе Word/LibreOffice могут подставить
    шрифт по умолчанию для русского текста."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), name)


def _write_cell(cell, text: str, *, align: str = "left", bold: bool = False,
                font_size: int = FONT_SIZE_PT) -> None:
    """Очистить ячейку и записать в неё текст с нужным выравниванием.
    Поддерживает '\\n' — даёт явный перенос строки внутри одной ячейки.
    """
    # Удалить все параграфы кроме первого.
    pars = cell.paragraphs
    for p in pars[1:]:
        p._element.getparent().remove(p._element)
    p = pars[0]
    # Удалить все старые runs первого параграфа.
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    parts = (text or "").split("\n")
    for i, part in enumerate(parts):
        if i > 0:
            br_run = p.add_run("")
            _apply_font(br_run, FONT_NAME, font_size, bold)
            br_run.add_break()
        run = p.add_run(part)
        _apply_font(run, FONT_NAME, font_size, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _cm_to_twips(cm: float) -> int:
    """1 inch = 1440 twips, 1 cm ≈ 567 twips. Точнее 567.0."""
    return int(round(cm * 567.0))


def _set_table_grid_and_layout(table, widths_cm: List[float]) -> None:
    """Жёстко прописать ширины колонок таблицы для LibreOffice/Word.

    Word + python-docx обычно довольствуется `cell.width` на каждую
    ячейку, но **LibreOffice** строже относится к разметке и при
    отсутствии явного `<w:tblGrid>` + `tblLayout type="fixed"` пересчитывает
    ширины по содержимому — длинная «Наименование работ/услуг»
    выдавливает крайние колонки, а наша колонка «Наименование» ужимается
    в 2 см. Без этих xml-настроек PDF едет.

    Делает три вещи:
      • выставляет `<w:tblW>` (общая ширина таблицы в dxa/twips)
      • переписывает `<w:tblGrid>` с явными `<w:gridCol w:w="…"/>`
      • выставляет `<w:tblLayout w:type="fixed"/>` — запрет авто-resize
    """
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    total_twips = sum(_cm_to_twips(w) for w in widths_cm)

    # 1) Полная ширина таблицы в twips.
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(total_twips))
    tblPr.append(tblW)

    # 2) Fixed layout — не давать LibreOffice пересчитывать ширины.
    for old in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    # 3) <w:tblGrid> — определение сетки колонок, перед строками.
    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w_cm in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(_cm_to_twips(w_cm)))
        grid.append(col)
    # Вставить сразу после tblPr (до первой <w:tr>).
    tblPr.addnext(grid)

    # 4) Продублировать на уровне ячеек — на случай если рендерер
    #    смотрит не на tblGrid а на tcW (Word так делает).
    table.autofit = False
    try:
        table.allow_autofit = False
    except Exception:
        pass
    for row in table.rows:
        for i, w_cm in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w_cm)


# Совместимость со старым именем — оставляем как тонкий wrapper.
def _set_col_widths(table, widths_cm: List[float]) -> None:
    _set_table_grid_and_layout(table, widths_cm)


# ────────────────────────────────────────────────────────────────────
# Построение таблицы товаров
# ────────────────────────────────────────────────────────────────────

def _build_products_table(doc, products: List[Dict[str, Any]]):
    """Создать docx-таблицу под образец и вернуть её.

    Колонки: №, Наименование, Ед. изм., Кол-во, Стоимость с НДС.
    Шапка — центр, жирный. Цифровые колонки — центр; «Наименование» —
    с двусторонним выравниванием (justify), как в исходнике.
    Итог — merge 0..3 + надпись «Итого с НДС, руб.:» справа жирным.
    Вставка в нужное место документа делается отдельно (см. main).
    """
    table = doc.add_table(rows=1 + len(products), cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 1) Шапка
    for i, h in enumerate(HEADERS):
        _write_cell(table.rows[0].cells[i], h, align="center", bold=True)

    # 2) Строки данных
    for r_idx, item in enumerate(products, start=1):
        row = table.rows[r_idx].cells
        _write_cell(row[0], str(item.get("idx", r_idx)), align="center")
        _write_cell(row[1], str(item.get("name", "")), align="justify")
        _write_cell(row[2], str(item.get("unit", "")), align="center")
        _write_cell(row[3], _fmt_qty(item.get("qty")), align="center")
        _write_cell(row[4], _fmt_money(item.get("total")), align="center")

    # 3) Итоговая строка (объединяем колонки 0..3 → подпись справа).
    total_sum = sum(float(it.get("total") or 0) for it in products)
    foot = table.add_row().cells
    merged = foot[0].merge(foot[1]).merge(foot[2]).merge(foot[3])
    _write_cell(merged, "Итого с НДС, руб.:", align="right", bold=True)
    _write_cell(foot[4], _fmt_money(total_sum), align="center", bold=True)

    # 4) Ширины колонок — в самом конце, когда вся структура готова.
    _set_col_widths(table, COL_WIDTHS_CM)

    return table


# ────────────────────────────────────────────────────────────────────
# Форматирование значений (под шаблон-образец)
# ────────────────────────────────────────────────────────────────────

def _fmt_money(v) -> str:
    """Деньги — точка как десятичный, неразрывные пробелы между тысячами:
    «30 900 000.00» (как в шаблоне-образце)."""
    try:
        n = float(v or 0)
    except (TypeError, ValueError):
        return "0.00"
    whole, frac = f"{n:,.2f}".split(".")
    whole = whole.replace(",", " ")  # NBSP вместо разделителя тысяч
    return f"{whole}.{frac}"


def _fmt_qty(v) -> str:
    try:
        n = float(v or 0)
    except (TypeError, ValueError):
        return "0"
    if abs(n - int(n)) < 1e-9:
        return str(int(n))
    return f"{n:g}"


# ────────────────────────────────────────────────────────────────────
# Вспомогательные XML-операции
# ────────────────────────────────────────────────────────────────────

def _move_table_before_paragraph(table, paragraph) -> None:
    """python-docx добавляет таблицу в конец. Переносим её на место
    параграфа-маркера через манипуляцию XML."""
    tbl = table._element
    p = paragraph._element
    p.addprevious(tbl)


def _remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


# ────────────────────────────────────────────────────────────────────
# Главная точка входа
# ────────────────────────────────────────────────────────────────────

def render_kp_docx(template_bytes: bytes, context: Dict[str, str], products: List[Dict[str, Any]]) -> bytes:
    """Главная точка входа.
      template_bytes  — содержимое исходного docx-шаблона
      context         — словарь плейсхолдер→значение (str-like)
      products        — список dict-ов: {idx, name, unit, qty, total}
    Возвращает bytes готового docx.
    """
    stream = io.BytesIO(template_bytes)
    doc = Document(stream)
    mapping = {str(k): ("" if v is None else str(v)) for k, v in context.items()}

    # 1) Сначала найдём параграф-маркер таблицы (если он есть).
    table_marker_para = None
    for p in list(doc.paragraphs):
        if "{products_table}" in p.text:
            table_marker_para = p
            break

    # 2) Заменяем плейсхолдеры в обычных параграфах (кроме маркерного —
    # его удалим целиком после вставки таблицы).
    for p in doc.paragraphs:
        if p is table_marker_para:
            continue
        _replace_in_paragraph(p, mapping)

    # 3) Заменяем плейсхолдеры в ячейках уже существующих таблиц
    # (реквизиты в шапке и т.п., если они вписаны в таблицу).
    for t in doc.tables:
        _replace_in_table(t, mapping)

    # 4) Вставляем настоящую таблицу товаров на место маркера.
    if table_marker_para is not None and products is not None:
        tbl = _build_products_table(doc, products)
        _move_table_before_paragraph(tbl, table_marker_para)
        _remove_paragraph(table_marker_para)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
