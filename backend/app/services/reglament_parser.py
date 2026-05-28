"""
Reglaments Phase 3: парсер PDF/DOCX → секции нормативного документа.

Цель — извлечь из загруженного файла структурированный список секций
вида `[{section_number, section_title, content, order_idx}]` для сохранения
в `reglament_sections`.

Эвристика разбиения:
  1. Идём по «абзацам»: для DOCX — `Document.paragraphs`, для PDF —
     извлечённый текст разбиваем по двойному \\n.
  2. Если абзац выглядит как заголовок раздела (паттерн `^N[.N]* Title`
     или DOCX-стиль Heading 1/2/3) — закрываем предыдущую секцию,
     начинаем новую.
  3. Если не нашли никаких заголовков — фолбэк: разбиение по чанкам
     ~500 слов (приоритет — лучше что-то, чем ничего).

Зависимости — `pypdf` и `python-docx`. Оба уже в системных deps.
Качество — best-effort; для красивого парсинга нужен ML (pdfplumber +
table detection + layout analysis), но это уже Phase 3+.

Public API:
    parse_pdf_bytes(blob)  → list[dict]
    parse_docx_bytes(blob) → list[dict]
    parse_file(blob, filename) → list[dict]   # автодиспетчер по расширению
"""
from __future__ import annotations

import io
import re
from typing import Iterable, List

# Паттерн заголовка раздела: «5.4.2 Защитный слой бетона», «1 Область
# применения», «Приложение А Условные обозначения».
# Допускаем nested-нумерацию (до 5 уровней) и спец-заголовки (Приложение,
# Введение, Содержание, Библиография).
_SECTION_RE = re.compile(
    r"^("
    r"\d+(?:\.\d+){0,5}"             # 5.4.2 / 1 / 10.3.1
    r"|Приложение\s+[А-ЯA-Z\d]+"     # Приложение А
    r"|Введение|Содержание|Библиография|Литература|Предисловие"
    r")\s+(.{3,200})$",
    re.MULTILINE,
)

# Эвристика «это заголовок» для DOCX-абзацев без явного стиля Heading.
# Короткий, заглавная первая буква, без точки в конце, нет жирного
# подсказки — оставим только проверку через _SECTION_RE.

_CHUNK_WORDS = 500     # фолбэк-чанк если не найдено заголовков
_MIN_SECTION_CHARS = 30  # секции короче — игнорируем (мусор)


def _detect_section_header(text: str) -> tuple[str | None, str | None]:
    """Если строка — заголовок раздела, возвращает (number, title).
    Иначе (None, None)."""
    s = (text or "").strip()
    if not s or len(s) > 240:
        return None, None
    m = _SECTION_RE.match(s)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None, None


def _normalize_paragraphs(paragraphs: Iterable[str]) -> List[str]:
    """Очистка: убираем колонтитулы, пустые строки, дубликаты подряд."""
    out: List[str] = []
    prev = None
    for p in paragraphs:
        p = (p or "").strip()
        if not p:
            continue
        # Колонтитулы вида «Стр. 5 из 158» — отсекаем
        if re.match(r"^(стр\.|страница|page)\s*\d+", p, re.I):
            continue
        if p == prev:
            continue
        out.append(p)
        prev = p
    return out


def _build_sections(paragraphs: List[str]) -> List[dict]:
    """Сборка секций из последовательности абзацев по заголовкам.
    Возвращает [{section_number, section_title, content, order_idx, char_count}]."""
    sections: List[dict] = []
    current = {"section_number": None, "section_title": "Введение", "content_parts": []}
    order = 0

    def close_current():
        nonlocal order
        body = "\n\n".join(current["content_parts"]).strip()
        if len(body) >= _MIN_SECTION_CHARS:
            sections.append({
                "section_number": current["section_number"] or "0",
                "section_title": current["section_title"] or "(без названия)",
                "content": body,
                "order_idx": order,
                "char_count": len(body),
            })
            order += 1

    for para in paragraphs:
        num, title = _detect_section_header(para)
        if num:
            close_current()
            current = {"section_number": num, "section_title": title, "content_parts": []}
        else:
            current["content_parts"].append(para)
    close_current()

    # Фолбэк: если найдено <2 секций — режем чанками по ~500 слов.
    if len(sections) < 2 and paragraphs:
        full_text = "\n\n".join(paragraphs)
        words = full_text.split()
        if len(words) < 50:
            return sections  # очень короткий документ — оставляем как есть
        chunked: List[dict] = []
        for i in range(0, len(words), _CHUNK_WORDS):
            chunk = " ".join(words[i : i + _CHUNK_WORDS])
            chunked.append({
                "section_number": str(len(chunked) + 1),
                "section_title": f"Часть {len(chunked) + 1}",
                "content": chunk,
                "order_idx": len(chunked),
                "char_count": len(chunk),
            })
        return chunked

    return sections


# ────────────────────────────────────────────────────────────────────
# Public API
# ────────────────────────────────────────────────────────────────────

def _extract_pdf_text_fitz(blob: bytes) -> str:
    """Извлечь текст через PyMuPDF (fitz). Намного качественнее pypdf
    для многоколоночного layout, формул, таблиц. Использует blocks-режим
    для сохранения структуры параграфов.
    """
    import fitz  # type: ignore
    doc = fitz.open(stream=blob, filetype="pdf")
    chunks: List[str] = []
    for page in doc:
        # `blocks` режим даёт лучший reading-order для многоколоночных
        # документов. Каждый block — это абзац/параграф.
        try:
            blocks = page.get_text("blocks")
            for b in blocks:
                # blocks возвращает (x0, y0, x1, y1, text, block_no, block_type)
                text = (b[4] or "").strip() if len(b) > 4 else ""
                if text:
                    chunks.append(text)
        except Exception:
            # Fallback на простой text mode
            try:
                t = page.get_text("text") or ""
                chunks.extend(p.strip() for p in t.split("\n\n") if p.strip())
            except Exception:
                pass
    doc.close()
    return "\n\n".join(chunks)


def _extract_pdf_text_pypdf(blob: bytes) -> str:
    """Fallback парсер на pypdf — если fitz недоступен."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(blob))
    chunks: List[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
            for c in re.split(r"\n\s*\n", text):
                c = c.replace("\n", " ").strip()
                if c:
                    chunks.append(c)
        except Exception:
            continue
    return "\n\n".join(chunks)


def parse_pdf_bytes(blob: bytes) -> List[dict]:
    """Распарсить PDF (bytes) → список секций.

    Приоритет: PyMuPDF (fitz) > pypdf. PyMuPDF существенно лучше работает
    с многоколоночным layout, формулами и таблицами строительной
    нормативки. Если fitz не установлен (dev-окружение) — fallback.
    """
    try:
        text = _extract_pdf_text_fitz(blob)
    except ImportError:
        text = _extract_pdf_text_pypdf(blob)
    except Exception:
        # Если fitz упал на конкретном PDF — ещё раз через pypdf
        try:
            text = _extract_pdf_text_pypdf(blob)
        except Exception:
            return []

    raw_paragraphs = re.split(r"\n\s*\n", text)
    return _build_sections(_normalize_paragraphs(raw_paragraphs))


def parse_docx_bytes(blob: bytes) -> List[dict]:
    """Распарсить DOCX (bytes) → список секций.

    Используем python-docx. У DOCX есть styled headings — это даёт более
    надёжную сегментацию, чем PDF. Если стиль Heading 1/2/3 — это явный
    заголовок (даже если номер не выделен).
    """
    from docx import Document

    doc = Document(io.BytesIO(blob))
    paragraphs: List[str] = []
    heading_markers: List[tuple[int, str]] = []  # (paragraph_index, style_name)

    for i, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if "heading" in style or style.startswith("заголовок"):
            heading_markers.append((len(paragraphs), text))
        paragraphs.append(text)

    paragraphs = _normalize_paragraphs(paragraphs)

    # Если есть styled headings — используем их как явные разделители.
    if heading_markers and len(heading_markers) >= 2:
        sections: List[dict] = []
        boundaries = [m[0] for m in heading_markers] + [len(paragraphs)]
        for idx, (start, end) in enumerate(zip(boundaries[:-1], boundaries[1:])):
            heading_text = paragraphs[start] if start < len(paragraphs) else ""
            num, title = _detect_section_header(heading_text)
            if not num:
                # Heading без явной нумерации — выставим автоматическую
                num = str(idx + 1)
                title = heading_text
            body_paras = paragraphs[start + 1 : end]
            body = "\n\n".join(body_paras).strip()
            if len(body) < _MIN_SECTION_CHARS:
                continue
            sections.append({
                "section_number": num,
                "section_title": title,
                "content": body,
                "order_idx": len(sections),
                "char_count": len(body),
            })
        if sections:
            return sections

    # Иначе — общая эвристика по тексту.
    return _build_sections(paragraphs)


def parse_file(blob: bytes, filename: str) -> List[dict]:
    """Авто-диспетчер по расширению файла."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return parse_pdf_bytes(blob)
    if name.endswith((".docx", ".doc")):
        # .doc (старый формат) python-docx не читает — но это редкий
        # случай для нормативки 2010+. Попробуем как DOCX, на ошибке
        # вернём пустой список.
        try:
            return parse_docx_bytes(blob)
        except Exception:
            return []
    # Неизвестный формат — пробуем как PDF (часто файлы без расширения).
    try:
        return parse_pdf_bytes(blob)
    except Exception:
        return []
